"""
DOCX Tools — Document Generator untuk Agent Office (Document Team / Ryo)
Membuat file .docx terformat untuk laporan, skripsi, proposal, dll.

Supported features:
- Custom font, size, line spacing
- Cover page (judul, subtitle, penulis, institusi, tahun)
- Heading hierarchy (H1-H3)
- Bold, italic, underline inline formatting
- Bullet & numbered lists
- Custom margins
- Page numbering
- Text alignment (left, center, right, justify)
- Table support
"""

import json
import os
import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re


# ============================================================
# PRESET FORMAT — Template siap pakai
# ============================================================

PRESETS = {
    "skripsi": {
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.5,
        "margins": {"top": 3, "bottom": 3, "left": 4, "right": 3},
        "alignment": "justify",
        "page_numbering": True
    },
    "laporan": {
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.5,
        "margins": {"top": 3, "bottom": 3, "left": 4, "right": 3},
        "alignment": "justify",
        "page_numbering": True
    },
    "proposal": {
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.5,
        "margins": {"top": 3, "bottom": 3, "left": 4, "right": 3},
        "alignment": "justify",
        "page_numbering": True
    },
    "surat": {
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.0,
        "margins": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
        "alignment": "justify",
        "page_numbering": False
    },
    "makalah": {
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.5,
        "margins": {"top": 3, "bottom": 3, "left": 4, "right": 3},
        "alignment": "justify",
        "page_numbering": True
    },
    "modern": {
        "font": "Calibri",
        "font_size": 11,
        "line_spacing": 1.15,
        "margins": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
        "alignment": "left",
        "page_numbering": True
    }
}

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
}


def _apply_run_format(run, font_name, font_size, bold=False, italic=False, underline=False, color=None):
    """Apply formatting to a text run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        if isinstance(color, str) and color.startswith("#"):
            r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
            run.font.color.rgb = RGBColor(r, g, b)
    # Force font name via XML (needed for East Asian font compatibility)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)


def _parse_inline_formatting(paragraph, text, font_name, font_size):
    """Parse inline markdown-like formatting: **bold**, *italic*, __underline__"""
    # Pattern: **bold**, *italic*, __underline__
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            _apply_run_format(run, font_name, font_size)
        
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            _apply_run_format(run, font_name, font_size, bold=True)
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            _apply_run_format(run, font_name, font_size, italic=True)
        elif match.group(4):  # __underline__
            run = paragraph.add_run(match.group(4))
            _apply_run_format(run, font_name, font_size, underline=True)
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        _apply_run_format(run, font_name, font_size)
    
    # If no formatting found, add plain text
    if last_end == 0 and text:
        run = paragraph.add_run(text)
        _apply_run_format(run, font_name, font_size)


def _add_page_number(doc):
    """Add page number at bottom center of each page."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = paragraph.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        
        run2 = paragraph.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        
        run3 = paragraph.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)


def _create_cover_page(doc, cover_data, font_name, font_size):
    """Create a formatted cover page."""
    # Add spacing at top
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    
    # Institution / Logo text
    if cover_data.get("institution"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_data["institution"].upper())
        _apply_run_format(run, font_name, 14, bold=True)
    
    # Faculty
    if cover_data.get("faculty"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_data["faculty"].upper())
        _apply_run_format(run, font_name, 14, bold=True)
    
    # Program Studi
    if cover_data.get("program"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_data["program"].upper())
        _apply_run_format(run, font_name, 14, bold=True)
    
    # Spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Document type (SKRIPSI, LAPORAN, etc.)
    if cover_data.get("doc_type"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_data["doc_type"].upper())
        _apply_run_format(run, font_name, 16, bold=True)
        p.paragraph_format.space_after = Pt(24)
    
    # Title
    if cover_data.get("title"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_data["title"].upper())
        _apply_run_format(run, font_name, 14, bold=True)
        p.paragraph_format.space_after = Pt(12)
    
    # Subtitle
    if cover_data.get("subtitle"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_data["subtitle"])
        _apply_run_format(run, font_name, 12)
    
    # Spacing before author
    for _ in range(4):
        doc.add_paragraph()
    
    # Disusun Oleh / Author info
    if cover_data.get("author"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Disusun Oleh:")
        _apply_run_format(run, font_name, 12)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_data["author"])
        _apply_run_format(run, font_name, 12, bold=True)
    
    # NIM
    if cover_data.get("nim"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"NIM: {cover_data['nim']}")
        _apply_run_format(run, font_name, 12)
    
    # Spacing
    for _ in range(4):
        doc.add_paragraph()
    
    # City and Year
    if cover_data.get("city") or cover_data.get("year"):
        city = cover_data.get("city", "")
        year = cover_data.get("year", str(datetime.datetime.now().year))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{city}\n{year}" if city else year)
        _apply_run_format(run, font_name, 12, bold=True)
    
    # Page break after cover
    doc.add_page_break()


def _add_table(doc, table_data, font_name, font_size):
    """Add a formatted table to the document."""
    headers = table_data.get("headers", [])
    rows_data = table_data.get("rows", [])
    
    if not headers and not rows_data:
        return
    
    num_cols = len(headers) if headers else len(rows_data[0]) if rows_data else 0
    num_rows = (1 if headers else 0) + len(rows_data)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    row_idx = 0
    
    # Add headers
    if headers:
        for col_idx, header_text in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(header_text))
            _apply_run_format(run, font_name, font_size, bold=True)
            # Header background color
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
            cell._element.get_or_add_tcPr().append(shading)
        row_idx = 1
    
    # Add data rows
    for row_data in rows_data:
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text))
                _apply_run_format(run, font_name, font_size)
        row_idx += 1
    
    # Add spacing after table
    doc.add_paragraph()


def create_docx(param_json: str) -> str:
    """
    Membuat file .docx berformat dari parameter JSON.
    
    Parameter JSON format:
    {
        "filename": "output.docx",
        "preset": "skripsi",           // optional: skripsi, laporan, proposal, surat, makalah, modern
        "settings": {                   // optional: override preset
            "font": "Times New Roman",
            "font_size": 12,
            "line_spacing": 1.5,
            "margins": {"top": 3, "bottom": 3, "left": 4, "right": 3},
            "alignment": "justify",
            "page_numbering": true
        },
        "cover": {                      // optional
            "institution": "Universitas XYZ",
            "faculty": "Fakultas Teknik",
            "program": "Program Studi Informatika",
            "doc_type": "SKRIPSI",
            "title": "Judul Skripsi",
            "subtitle": "Diajukan untuk memenuhi...",
            "author": "Nama Lengkap",
            "nim": "12345678",
            "city": "Jakarta",
            "year": "2026"
        },
        "content": [
            {"type": "heading1", "text": "BAB I PENDAHULUAN"},
            {"type": "heading2", "text": "1.1 Latar Belakang"},
            {"type": "paragraph", "text": "Isi paragraf..."},
            {"type": "bullet_list", "items": ["item 1", "item 2"]},
            {"type": "numbered_list", "items": ["item 1", "item 2"]},
            {"type": "table", "headers": ["No", "Nama"], "rows": [["1", "Test"]]},
            {"type": "page_break"},
            {"type": "heading1", "text": "BAB II TINJAUAN PUSTAKA"}
        ]
    }
    """
    try:
        data = json.loads(param_json)
    except json.JSONDecodeError as e:
        return f"[ERROR] JSON tidak valid: {e}"
    
    filename = data.get("filename", "document.docx")
    if not filename.endswith(".docx"):
        filename += ".docx"
    
    # Resolve output path
    output_dir = os.path.join(".", "output_docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    
    # Load preset or default settings
    preset_name = data.get("preset", "laporan")
    settings = PRESETS.get(preset_name, PRESETS["laporan"]).copy()
    
    # Override with custom settings if provided
    if "settings" in data:
        custom = data["settings"]
        settings["font"] = custom.get("font", settings["font"])
        settings["font_size"] = custom.get("font_size", settings["font_size"])
        settings["line_spacing"] = custom.get("line_spacing", settings["line_spacing"])
        settings["alignment"] = custom.get("alignment", settings["alignment"])
        settings["page_numbering"] = custom.get("page_numbering", settings["page_numbering"])
        if "margins" in custom:
            settings["margins"].update(custom["margins"])
    
    font_name = settings["font"]
    font_size = settings["font_size"]
    line_spacing = settings["line_spacing"]
    alignment = ALIGNMENT_MAP.get(settings["alignment"], WD_ALIGN_PARAGRAPH.JUSTIFY)
    margins = settings["margins"]
    
    # Create document
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(margins.get("top", 3))
        section.bottom_margin = Cm(margins.get("bottom", 3))
        section.left_margin = Cm(margins.get("left", 4))
        section.right_margin = Cm(margins.get("right", 3))
    
    # Set default font for Normal style
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.alignment = alignment
    # Force font via XML
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.insert(0, rFonts)
    
    # Configure heading styles
    for level in range(1, 4):
        heading_style_name = f'Heading {level}'
        if heading_style_name in doc.styles:
            h_style = doc.styles[heading_style_name]
            h_style.font.name = font_name
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            h_sizes = {1: font_size + 2, 2: font_size + 1, 3: font_size}
            h_style.font.size = Pt(h_sizes.get(level, font_size))
            h_rPr = h_style.element.get_or_add_rPr()
            h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
            h_rPr.insert(0, h_rFonts)
    
    # Create cover page if provided
    if "cover" in data:
        _create_cover_page(doc, data["cover"], font_name, font_size)
    
    # Remove default empty paragraph
    if doc.paragraphs and doc.paragraphs[0].text == '' and "cover" not in data:
        p_element = doc.paragraphs[0]._element
        p_element.getparent().remove(p_element)
    
    # Process content blocks
    content = data.get("content", [])
    for block in content:
        block_type = block.get("type", "paragraph")
        text = block.get("text", "")
        
        if block_type == "heading1":
            p = doc.add_heading(text, level=1)
            p.alignment = alignment
            for run in p.runs:
                _apply_run_format(run, font_name, font_size + 2, bold=True)
        
        elif block_type == "heading2":
            p = doc.add_heading(text, level=2)
            p.alignment = alignment
            for run in p.runs:
                _apply_run_format(run, font_name, font_size + 1, bold=True)
        
        elif block_type == "heading3":
            p = doc.add_heading(text, level=3)
            p.alignment = alignment
            for run in p.runs:
                _apply_run_format(run, font_name, font_size, bold=True)
        
        elif block_type == "paragraph":
            p = doc.add_paragraph()
            p.alignment = alignment
            p.paragraph_format.line_spacing = line_spacing
            # First line indent for body text (common in skripsi)
            if preset_name in ("skripsi", "laporan", "proposal", "makalah"):
                p.paragraph_format.first_line_indent = Cm(1.27)  # ~0.5 inch
            _parse_inline_formatting(p, text, font_name, font_size)
        
        elif block_type == "bullet_list":
            items = block.get("items", [])
            for item in items:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.line_spacing = line_spacing
                _parse_inline_formatting(p, item, font_name, font_size)
        
        elif block_type == "numbered_list":
            items = block.get("items", [])
            for item in items:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.line_spacing = line_spacing
                _parse_inline_formatting(p, item, font_name, font_size)
        
        elif block_type == "table":
            _add_table(doc, block, font_name, font_size)
        
        elif block_type == "page_break":
            doc.add_page_break()
        
        elif block_type == "empty_line":
            count = block.get("count", 1)
            for _ in range(count):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
        
        elif block_type == "image":
            img_path = block.get("path", "")
            width_cm = block.get("width_cm", 15)
            
            if os.path.exists(img_path):
                # Add image paragraph
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Cm(width_cm))
                
                # Add optional caption
                caption = block.get("caption", "")
                if caption:
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap = cap_p.add_run(caption)
                    _apply_run_format(run_cap, font_name, font_size - 2, italic=True)
            else:
                p = doc.add_paragraph(f"[ERROR: Gambar tidak ditemukan di {img_path}]")
                _apply_run_format(p.add_run(), font_name, font_size, italic=True)
    
    # Add page numbering if enabled
    if settings.get("page_numbering"):
        _add_page_number(doc)
    
    # Save document
    doc.save(output_path)
    abs_path = os.path.abspath(output_path)
    
    return (
        f"[SUCCESS] Dokumen berhasil dibuat!\n"
        f"📄 File: {abs_path}\n"
        f"📐 Format: {font_name} {font_size}pt, spasi {line_spacing}\n"
        f"📏 Margin: L={margins['left']}cm T={margins['top']}cm R={margins['right']}cm B={margins['bottom']}cm\n"
        f"📑 Halaman: {'Bernomor' if settings.get('page_numbering') else 'Tanpa nomor'}\n"
        f"📝 Content blocks: {len(content)}"
    )


def list_presets(param: str = "") -> str:
    """List semua preset format dokumen yang tersedia."""
    output = "📋 Preset Format Dokumen:\n\n"
    for name, settings in PRESETS.items():
        output += f"**{name}**\n"
        output += f"  Font: {settings['font']} {settings['font_size']}pt\n"
        output += f"  Spasi: {settings['line_spacing']}\n"
        output += f"  Margin: L={settings['margins']['left']}cm T={settings['margins']['top']}cm R={settings['margins']['right']}cm B={settings['margins']['bottom']}cm\n"
        output += f"  Alignment: {settings['alignment']}\n\n"
    return output


def convert_to_pdf(docx_path: str) -> str:
    """Konversi file .docx ke .pdf menggunakan docx2pdf (butuh MS Word di Windows)."""
    if not os.path.exists(docx_path):
        return f"[ERROR] File tidak ditemukan: {docx_path}"
    
    try:
        from docx2pdf import convert
        import pythoncom
        
        # Inisialisasi COM untuk multithreading (FastAPI/Thread safety)
        pythoncom.CoInitialize()
        
        pdf_path = docx_path.replace(".docx", ".pdf")
        print(f"[DOCX2PDF] Converting {docx_path} to {pdf_path}...")
        
        convert(docx_path, pdf_path)
        
        abs_path = os.path.abspath(pdf_path)
        return f"[SUCCESS] PDF berhasil dibuat!\n📄 File: {abs_path}"
    except Exception as e:
        print(f"[DOCX2PDF ERROR] {e}")
        return f"[ERROR] Gagal konversi ke PDF: {e}\n(Pastikan Microsoft Word terinstal dan file tidak sedang dibuka)"
    finally:
        try:
            pythoncom.CoUninitialize()
        except:
            pass
