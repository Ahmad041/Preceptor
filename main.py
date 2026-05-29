from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request as FastAPIRequest
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import requests
import datetime
try:
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    GOOGLE_AUTH_AVAILABLE = True
except ImportError:
    GOOGLE_AUTH_AVAILABLE = False
    Credentials = None
import base64
import hashlib
import os
import asyncio
import io
import re
import numpy as np
import json
import asyncio
from datetime import datetime as dt
from dotenv import load_dotenv
import soundfile as sf

try:
    import google.generativeai as genai
    GENAI_AVAILABLE = True
except ImportError:
    GENAI_AVAILABLE = False
    genai = None

from PIL import ImageGrab
import psutil
import os_tools
from agent_tools import process_agent_command_with_tools
import agent_logger
from memory_system import memory
from notes_engine import notes_index, build_note_metadata, get_watched_folders, add_watched_folder, remove_watched_folder
from embedding_engine import embedding_engine
from desktop_pilot import (
    desktop_click,
    desktop_type,
    desktop_press,
    desktop_hotkey,
    desktop_scroll,
    desktop_screenshot,
    get_screen_info
)
import desktop_pilot
from vision_loop import vision_engine
from research_loop import research_engine
from jarvis_orchestrator import jarvis
from gitnexus_runner import gitnexus_server
from omniscient import omniscient
import tech_doc_generator as tech_gen
# MCP Support
try:
    from mcp.server.sse import SseServerTransport
    from mcp_server import bocchi_mcp_server
    from mcp_client import mcp_registry
except ImportError:
    SseServerTransport = None
    bocchi_mcp_server = None
    mcp_registry = None
    print("[WARNING] MCP module not found, MCP features will be disabled.")

load_dotenv()
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "qwen/qwen-2.5-72b-instruct")

if not OPENROUTER_API_KEY:
    print("[WARNING] OPENROUTER_API_KEY belum di-set di file .env!")

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
if GEMINI_API_KEY:
    genai.configure(
        api_key=GEMINI_API_KEY,
        client_options={"api_endpoint": "generativelanguage.googleapis.com"}
    )
else:
    print("[WARNING] GEMINI_API_KEY belum di-set di file .env!")

# --- Library untuk membaca dokumen ---
import PyPDF2
from docx import Document as DocxDocument
import docx2pdf

# --- SURAT IZIN KHUSUS UNTUK KOMPOR PYTORCH 2.6+ ---
import torch
_original_load = torch.load
def _patched_load(*args, **kwargs):
    if 'weights_only' not in kwargs:
        kwargs['weights_only'] = False
    return _original_load(*args, **kwargs)
torch.load = _patched_load
# ---------------------------------------------------

import sys
import io

# ── FIX: Windows charmap encoding error ──────────────────────
# Windows default console encoding (cp1252) can't handle emoji/Unicode.
# Force UTF-8 for all print() output.
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import re
def sanitize_for_tts(text):
    """Strip emoji and non-BMP Unicode characters that crash TTS/encoding."""
    # Remove emoji and symbols (Unicode blocks: Emoticons, Symbols, Dingbats, etc.)
    text = re.sub(r'[\U0001F600-\U0001F64F]', '', text)  # Emoticons
    text = re.sub(r'[\U0001F300-\U0001F5FF]', '', text)  # Misc Symbols
    text = re.sub(r'[\U0001F680-\U0001F6FF]', '', text)  # Transport
    text = re.sub(r'[\U0001F1E0-\U0001F1FF]', '', text)  # Flags
    text = re.sub(r'[\U00002702-\U000027B0]', '', text)  # Dingbats
    text = re.sub(r'[\U0000FE00-\U0000FE0F]', '', text)  # Variation selectors
    text = re.sub(r'[\U0000200D]', '', text)              # ZWJ
    text = re.sub(r'[\U000023E9-\U000023F3]', '', text)   # Misc technical (⏳ etc)
    text = re.sub(r'[\U00002600-\U000026FF]', '', text)   # Misc symbols
    text = re.sub(r'[\U0000FE00-\U0000FEFF]', '', text)   # Specials
    text = re.sub(r'[\U00010000-\U0010FFFF]', '', text)   # All supplementary planes
    # Clean up extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "Qwen3-TTS"))

app = FastAPI()

# Setup direktori cache
AUDIO_CACHE_DIR = os.path.join(os.getcwd(), "data", "audio_cache")
os.makedirs(AUDIO_CACHE_DIR, exist_ok=True)

# ============================================================
# GLOBAL SETTINGS (HOT-RELOADABLE)
# ============================================================
GLOBAL_SETTINGS = {
    "language": "ID",
    "llm_engine": "Ollama",
    "llm_model": "llama3",
    "tts_model_chat": "Qwen/Qwen3-TTS-12Hz-0.6B-Base",
    "tts_model_story": "Qwen/Qwen3-TTS-12Hz-1.7B-Base",
    "vram_gb": 0,
    "system_prompt": "Kamu adalah Hitori Gotou (Bocchi). Kamu pemalu, suka musik, dan sedikit gugup saat bicara.",
    "user_nama": "Senpai",
    "user_hubungan": "Teman",
    "visual_mode": "2D", # 2D or 3D
    "volume": 1.0,
    "is_muted": False
}

def load_global_settings():
    try:
        if os.path.exists("spec.json"):
            with open("spec.json", 'r', encoding='utf-8') as f:
                saved = json.load(f)
                GLOBAL_SETTINGS.update(saved)
    except Exception as e:
        print(f"[WARNING] Gagal memuat GLOBAL_SETTINGS: {e}")

def save_global_settings():
    try:
        with open("spec.json", 'w', encoding='utf-8') as f:
            json.dump(GLOBAL_SETTINGS, f, indent=4)
    except Exception as e:
        print(f"[WARNING] Gagal menyimpan GLOBAL_SETTINGS: {e}")

load_global_settings()

# Global store untuk Retrieval-Augmented Generation (RAG)
rag_store = []

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================
# INITIALIZE NOTES & EMBEDDING ENGINE
# ============================================================
@app.on_event("startup")
async def startup_event():
    print("\n[SISTEM] Initializing Company Mode...")
    
    # Start GitNexus Server
    gitnexus_server.start()
    
    # 1. Load notes index from cache or scan
    if not notes_index.load_cache():
        notes_index.scan_all()
    
    # 2. Background embedding generation (so startup is not blocked)
    asyncio.create_task(initialize_embeddings())
    
    # 3. Start Research Engine (AI Co-Scientist)
    try:
        research_engine.start()
    except Exception as e:
        print(f"[WARNING] Gagal start Research Engine: {e}")

@app.on_event("shutdown")
async def shutdown_event():
    print("\n[SISTEM] Shutting down...")
    gitnexus_server.stop()

async def initialize_embeddings():
    print("[SISTEM] Generating/Updating Note Embeddings...")
    try:
        embedding_engine.embed_notes(notes_index)
        print("[SISTEM] [OK] Note Embeddings ready!")
    except Exception as e:
        print(f"[WARNING] Gagal generate embeddings: {e}")

# ============================================================
# 3. JARVIS ORCHESTRATOR
# ============================================================
class NoteCreate(BaseModel):
    title: str
    content: str = ""
    folder: Optional[str] = None
    tags: Optional[List[str]] = None

class NoteUpdate(BaseModel):
    content: str

class NoteAsk(BaseModel):
    question: str
    note_id: Optional[str] = None

class DeepSearchRequest(BaseModel):
    query: str
    include_web: bool = True

class UnifiedSearchRequest(BaseModel):
    query: str


# ============================================================
# OMNISCIENT (Unified Knowledge Hub)
# ============================================================
@app.post("/api/knowledge/search")
async def unified_knowledge_search(req: UnifiedSearchRequest):
    try:
        results = omniscient.unified_search(req.query)
        return {"status": "success", "data": results}
    except Exception as e:
        return {"status": "error", "message": str(e)}

# ============================================================
# 1. INISIALISASI QWEN3-TTS (Voice Cloning Mode) - DYNAMIC SWAPPING
CURRENT_TTS_MODEL_NAME = None
QWEN_TTS_MODEL = None
QWEN_TTS_TOKENIZER = None

def get_tts_model(mode="chat"):
    global QWEN_TTS_MODEL, CURRENT_TTS_MODEL_NAME
    
    target_model_key = "tts_model_chat" if mode == "chat" else "tts_model_story"
    target_model = GLOBAL_SETTINGS.get(target_model_key, "Qwen/Qwen3-TTS-12Hz-0.6B-Base")
    
    import torch
    device_tts = "cuda" if torch.cuda.is_available() else "cpu"
    if device_tts == "cuda":
        GLOBAL_SETTINGS["vram_gb"] = torch.cuda.get_device_properties(0).total_memory / (1024**3)
    
    if QWEN_TTS_MODEL is not None and CURRENT_TTS_MODEL_NAME == target_model:
        return QWEN_TTS_MODEL
        
    print(f"\n[SISTEM] TTS Swap: Memuat model untuk {mode.upper()} mode -> {target_model}")
    
    if QWEN_TTS_MODEL is not None:
        print(f"[SISTEM] Mengosongkan VRAM (Unloading {CURRENT_TTS_MODEL_NAME})...")
        del QWEN_TTS_MODEL
        import gc
        gc.collect()
        if device_tts == "cuda":
            torch.cuda.empty_cache()
            
    print(f"[SISTEM] Memuat {target_model} ke {device_tts.upper()}...")
    try:
        from qwen_tts import Qwen3TTSModel
        try:
            QWEN_TTS_MODEL = Qwen3TTSModel.from_pretrained(
                target_model,
                device_map=device_tts,
                torch_dtype=torch.bfloat16,
                attn_implementation="sdpa",
            )
            print(f"[SISTEM] [OK] Qwen3-TTS berhasil dimuat dengan SDPA di {device_tts.upper()}!")
        except Exception as e:
            print(f"[WARNING] Gagal memuat Qwen3-TTS dengan SDPA: {e}")
            print("[SISTEM] Mencoba tanpa SDPA...")
            QWEN_TTS_MODEL = Qwen3TTSModel.from_pretrained(
                target_model,
                device_map=device_tts,
                torch_dtype=torch.bfloat16,
            )
            print(f"[SISTEM] [OK] Qwen3-TTS berhasil dimuat di {device_tts.upper()} (tanpa SDPA)!")
            
        CURRENT_TTS_MODEL_NAME = target_model
    except Exception as e2:
        print(f"[WARNING] Gagal memuat Qwen3-TTS: {e2}")
        QWEN_TTS_MODEL = None
        CURRENT_TTS_MODEL_NAME = None
        
    return QWEN_TTS_MODEL

# Mapping emosi → instruksi suara Qwen3-TTS
EMOSI_INSTRUKSI = {
    "Joy":      "Speak in a cheerful, bright, and slightly excited tone. Voice should sound warm and happy.",
    "Angry":    "Speak in a frustrated, tense tone with slightly higher pitch. Sound irritated but still shy.",
    "Sorrow":   "Speak in a soft, melancholic tone. Voice should sound sad and a bit trembling.",
    "Fun":      "Speak in a playful, energetic tone. Sound like having fun and giggling.",
    "Surprised": "Speak in a surprised, slightly flustered tone. Voice cracks a little from shock.",
    "Neutral":  "Speak in a calm, quiet, and slightly nervous tone. Sound gentle and introverted."
}

# ============================================================
# 2. LOCALDOCS + RAG — Semantic Search dengan Nomic Embed
# ============================================================
# Daftar file yang sudah di-upload (untuk UI)
file_registry: List[dict] = []

GEMINI_EMBED_MODEL = "text-embedding-004"
CHUNK_SIZE = 500    # Jumlah karakter per chunk
CHUNK_OVERLAP = 50  # Overlap antar chunk agar konteks tidak terpotong
TOP_K = 5           # Jumlah chunk paling relevan yang dikirim ke OpenAI

def ekstrak_teks_dari_file(nama_file: str, konten: bytes) -> str:
    """Mengekstrak teks dari berbagai jenis file"""
    ekstensi = nama_file.lower().split('.')[-1]
    
    if ekstensi == 'txt':
        return konten.decode('utf-8', errors='ignore')
    elif ekstensi == 'pdf':
        reader = PyPDF2.PdfReader(io.BytesIO(konten))
        teks = ""
        for halaman in reader.pages:
            teks += halaman.extract_text() or ""
        return teks
    elif ekstensi in ['docx', 'doc']:
        doc = DocxDocument(io.BytesIO(konten))
        return "\n".join([p.text for p in doc.paragraphs])
    elif ekstensi in ['md', 'csv', 'json', 'py', 'js', 'html', 'css', 'jsx', 'ts', 'tsx']:
        return konten.decode('utf-8', errors='ignore')
    else:
        return f"[Format .{ekstensi} belum didukung]"

def potong_teks_jadi_chunk(teks: str, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP) -> List[str]:
    """Memotong teks panjang menjadi potongan-potongan kecil (chunk)"""
    chunks = []
    start = 0
    while start < len(teks):
        end = start + chunk_size
        chunk = teks[start:end].strip()
        if chunk:  # Abaikan chunk kosong
            chunks.append(chunk)
        start = end - overlap  # Mundur sedikit untuk overlap
    return chunks


# ============================================================
# ENDPOINTS
# ============================================================

@app.post("/api/upload")
async def upload_dokumen(file: UploadFile = File(...)):
    """Upload dokumen → ekstrak teks → chunk → embed → simpan"""
    try:
        konten = await file.read()
        teks = ekstrak_teks_dari_file(file.filename, konten)
        
        if not teks.strip():
            return {"status": "gagal", "error": "Dokumen kosong atau tidak bisa dibaca"}
        
        # Hapus chunks lama dari file yang sama (jika re-upload)
        memory.long_term_memory = [c for c in memory.long_term_memory if c.get("nama") != file.filename]
        
        # Potong teks jadi chunks
        chunks = potong_teks_jadi_chunk(teks)
        print(f"\n[DOKUMEN] '{file.filename}' → {len(chunks)} chunks")
        
        # Buat embedding untuk semua chunks sekaligus (batch)
        print(f"[EMBED] Membuat embedding untuk {len(chunks)} chunks dengan Nomic API...")
        embeddings = memory.create_embedding(chunks)
        
        if len(embeddings) != len(chunks):
            print(f"[WARNING] Jumlah embedding ({len(embeddings)}) != chunks ({len(chunks)})")
            # Fallback: hanya proses yang berhasil
            embeddings = embeddings[:len(chunks)]
        
        # Simpan ke RAG store
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            memory.add_to_long_term_memory(file.filename, chunk, emb)
        
        # Update file registry
        global file_registry
        file_registry = [f for f in file_registry if f["nama"] != file.filename]
        file_registry.append({
            "nama": file.filename,
            "panjang": len(teks),
            "chunks": len(chunks)
        })
        
        print(f"[DOKUMEN] ✅ '{file.filename}' berhasil diproses! ({len(chunks)} chunks, {len(teks)} karakter)")
        
        return {
            "status": "berhasil",
            "nama": file.filename,
            "panjang": len(teks),
            "chunks": len(chunks),
            "total_dokumen": len(file_registry)
        }
    except Exception as e:
        print(f"[ERROR] Gagal memproses dokumen: {e}")
        return {"status": "gagal", "error": str(e)}

@app.get("/api/dokumen")
async def daftar_dokumen():
    """Melihat daftar dokumen yang sudah di-upload"""
    return {
        "total": len(file_registry),
        "total_chunks": len(memory.long_term_memory),
        "dokumen": file_registry
    }

@app.delete("/api/dokumen/{nama_file}")
async def hapus_dokumen(nama_file: str):
    """Menghapus dokumen dari RAG store"""
    global file_registry
    sebelum = len(memory.long_term_memory)
    memory.long_term_memory = [c for c in memory.long_term_memory if c.get("nama") != nama_file]
    file_registry = [f for f in file_registry if f["nama"] != nama_file]
    
    dihapus = sebelum - len(memory.long_term_memory)
    if dihapus > 0:
        print(f"[DOKUMEN] Dihapus: {nama_file} ({dihapus} chunks)")
        return {"status": "dihapus", "chunks_dihapus": dihapus, "total_dokumen": len(file_registry)}
    return {"status": "tidak ditemukan"}

# ============================================================
# 3. CHAT ENDPOINT — Sekarang dengan RAG
# ============================================================
class PesanMasuk(BaseModel):
    pesan: str
    user_nama: Optional[str] = "Senpai"
    user_hubungan: Optional[str] = "Teman"
    lihat_layar: Optional[bool] = False

class ExecuteToolRequest(BaseModel):
    tool: str
    parameter: str
    pesan_asli: str
    izin_diberikan: bool

# Model untuk membuat file
class BuatFileRequest(BaseModel):
    nama: str
    konten: str

@app.post("/api/buat-file")
async def buat_file_api(data: BuatFileRequest):
    """Membuat file .txt atau .md di folder 'catatan'"""
    try:
        # Pastikan folder 'catatan' ada
        folder = "catatan"
        if not os.path.exists(folder):
            os.makedirs(folder)
            
        # Bersihkan nama file agar aman
        nama_aman = "".join([c for c in data.nama if c.isalnum() or c in "._- "]).strip()
        if not nama_aman:
            return {"status": "gagal", "error": "Nama file tidak valid"}
            
        # Tambahkan ekstensi jika belum ada
        if not (nama_aman.endswith(".txt") or nama_aman.endswith(".md")):
            nama_aman += ".txt"
            
        path_file = os.path.join(folder, nama_aman)
        
        with open(path_file, "w", encoding="utf-8") as f:
            f.write(data.konten)
            
        print(f"[FILE] ✅ Berhasil membuat file: {path_file}")
        return {"status": "berhasil", "path": path_file}
    except Exception as e:
        print(f"[FILE ERROR] Gagal membuat file: {e}")
        return {"status": "gagal", "error": str(e)}

@app.get("/api/list-catatan")
async def list_catatan():
    """Melihat daftar file di folder 'catatan'"""
    try:
        folder = "catatan"
        if not os.path.exists(folder):
            return {"catatan": []}
        
        files = []
        for f in os.listdir(folder):
            if os.path.isfile(os.path.join(folder, f)):
                stats = os.stat(os.path.join(folder, f))
                files.append({
                    "nama": f,
                    "ukuran": stats.st_size,
                    "waktu": stats.st_mtime
                })
        # Urutkan berdasarkan waktu terbaru
        files.sort(key=lambda x: x["waktu"], reverse=True)
        return {"catatan": files}
    except Exception as e:
        return {"status": "gagal", "error": str(e)}

@app.delete("/api/catatan/{nama_file}")
async def hapus_catatan(nama_file: str):
    """Menghapus file di folder 'catatan'"""
    try:
        path = os.path.join("catatan", nama_file)
        if os.path.exists(path):
            os.remove(path)
            return {"status": "berhasil"}
        return {"status": "gagal", "error": "File tidak ditemukan"}
    except Exception as e:
        return {"status": "gagal", "error": str(e)}

@app.post("/api/execute-tool")
async def execute_tool_api(data: ExecuteToolRequest):
    if not data.izin_diberikan:
        # Jika user menolak dari UI, kita sampaikan ke Bocchi bahwa user menolak
        data_chat = PesanMasuk(pesan=f"{data.pesan_asli}\n\n[SISTEM] Akses tool {data.tool} ditolak oleh Senpai. Minta maaf dan respon natural.")
        return await chat_dengan_ai(data_chat)
    
    # Eksekusi tool
    if data.tool in os_tools.SAFE_TOOLS:
        hasil_tool = os_tools.SAFE_TOOLS[data.tool](data.parameter)
        # Sampaikan ke Bocchi hasil dari tool
        data_chat = PesanMasuk(pesan=f"{data.pesan_asli}\n\n[SISTEM] Kamu baru saja mengeksekusi alat komputer: {data.tool}. Hasilnya adalah:\n{hasil_tool}\nRespons natural ke user (jangan tunjukkan format tool, cukup ngobrol).")
        
        response_ai = await chat_dengan_ai(data_chat)
        
        # Tambahkan ke canvas jika relevan
        if data.tool in ["baca_halaman_web", "cari_di_internet", "baca_file"]:
            response_ai["canvas_content"] = hasil_tool
            
        return response_ai
        
    return {"status": "gagal", "error": "Tool tidak valid"}

# ============================================================
# AGENT MISSION CONTROL ENDPOINT
# ============================================================

class AgentCommand(BaseModel):
    agent_id: str
    command: str
    conversation: Optional[List[dict]] = None
    focus_modes: Optional[List[str]] = None

@app.post("/api/agent/command")
async def agent_command_api(data: AgentCommand):
    """Memproses perintah untuk agent tertentu berdasarkan persona — dengan Tool Calling"""
    # Simpan focus modes yang aktif
    agent_logger.set_agent_focus_modes(data.agent_id, data.focus_modes)
    try:
        # === INTERCEPT ORCHESTRATOR ===
        if data.agent_id == "orchestrator":
            from orchestrator import orchestrator
            agent_logger.set_agent_status("orchestrator", "processing")
            agent_logger.record_command("orchestrator")
            agent_logger.log_activity("orchestrator", f"Planning flow: {data.command}", "info")
            
            flow = orchestrator.start_flow(data.command)
            tasks = orchestrator.plan_tasks(data.command)
            
            headers = {
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "HTTP-Referer": "http://localhost:3000",
                "X-Title": "Agent Office Mission Control",
                "Content-Type": "application/json"
            }
            
            overall_response = f"**PentAGI Flow Started**\nObjective: {data.command}\nPlanned Tasks: {len(tasks)}\n\n"
            
            for idx, task in enumerate(tasks, 1):
                target_agent = orchestrator.assign_agent(task.assignee_role)
                agent_logger.log_activity("orchestrator", f"Task {idx}: {task.description} -> {target_agent}", "info")
                
                # Fetch target persona
                target_persona_path = os.path.join("personas", f"{target_agent}.md")
                if os.path.exists(target_persona_path):
                    with open(target_persona_path, "r", encoding="utf-8") as f:
                        target_persona = f.read()
                else:
                    target_persona = "You are a helpful AI."
                
                # Context
                agent_context = orchestrator.get_agent_context(target_agent)
                target_persona += f"\n\n[ORCHESTRATOR INSTRUCTIONS]\nYou have been assigned the following task:\nTask: {task.description}\n{agent_context}"
                
                target_messages = [{"role": "system", "content": ""}, {"role": "user", "content": f"Please execute this task: {task.description}"}]
                
                agent_logger.set_agent_status(target_agent, "processing")
                
                evolve_chain = None
                if target_agent == "evolve":
                    from agent_tools import EVOLVE_MODEL_CHAIN
                    evolve_chain = EVOLVE_MODEL_CHAIN
                
                try:
                    task_response = await process_agent_command_with_tools(
                        persona_content=target_persona,
                        messages=target_messages,
                        headers=headers,
                        model=OPENROUTER_MODEL,
                        max_tool_rounds=7 if target_agent == "evolve" else 3,
                        agent_id=target_agent,
                        model_chain=evolve_chain
                    )
                    task.status = "completed"
                    task.result = task_response
                    overall_response += f"### Task {idx}: {task.description}\n**Agent**: {target_agent}\n**Result**:\n{task_response}\n\n"
                    
                    # Store memory
                    memory.record_episode(target_agent, task.description, "execute_task", task_response, True)
                    agent_logger.set_agent_status(target_agent, "done")
                except Exception as e:
                    task.status = "failed"
                    task.result = str(e)
                    overall_response += f"### Task {idx}: {task.description}\n**Agent**: {target_agent}\n**Result**: FAILED ({str(e)})\n\n"
                    memory.record_episode(target_agent, task.description, "execute_task", str(e), False)
                    agent_logger.set_agent_status(target_agent, "error")
                    
            flow.status = "completed"
            agent_logger.set_agent_status("orchestrator", "done")
            
            return {
                "status": "berhasil",
                "agent_id": "orchestrator",
                "response": overall_response
            }

        # === NORMAL AGENT EXECUTION ===
        persona_path = os.path.join("personas", f"{data.agent_id}.md")
        if not os.path.exists(persona_path):
            return {"status": "gagal", "error": f"Persona for {data.agent_id} not found"}
            
        with open(persona_path, "r", encoding="utf-8") as f:
            persona_content = f.read()
        
        # === REAL-TIME LOGGING ===
        agent_logger.set_agent_status(data.agent_id, "processing")
        agent_logger.record_command(data.agent_id)
        cmd_preview = data.command[:60] if data.command else "(conversation)"
        agent_logger.log_activity(data.agent_id, f"Received: {cmd_preview}", "info")
        
        # Build messages — gunakan conversation history jika ada
        messages = [{"role": "system", "content": ""}]  # placeholder, akan di-replace oleh agent_tools
        if data.conversation and len(data.conversation) > 0:
            messages.extend(data.conversation)
        else:
            messages.append({"role": "user", "content": data.command})
        
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "HTTP-Referer": "http://localhost:3000",
            "X-Title": "Agent Office Mission Control",
            "Content-Type": "application/json"
        }
        
        agent_logger.log_activity(data.agent_id, "Connecting to AI model...", "system")
        
        # Evolve agent butuh lebih banyak putaran untuk siklus evolusi penuh
        evolve_rounds = 7 if data.agent_id == "evolve" else 3
        
        # Import model chain untuk agen Evolve (multi-model fallback)
        evolve_chain = None
        if data.agent_id == "evolve":
            from agent_tools import EVOLVE_MODEL_CHAIN
            evolve_chain = EVOLVE_MODEL_CHAIN
            agent_logger.log_activity(
                data.agent_id,
                f"AlphaEvolve mode: {len(evolve_chain)} model candidates loaded",
                "system"
            )
        
        # Gunakan tool-calling loop dari agent_tools
        ai_response = await process_agent_command_with_tools(
            persona_content=persona_content,
            messages=messages,
            headers=headers,
            model=OPENROUTER_MODEL,
            max_tool_rounds=evolve_rounds,
            agent_id=data.agent_id,
            model_chain=evolve_chain
        )
        
        agent_logger.log_activity(data.agent_id, "Response generated OK", "success")
        agent_logger.set_agent_status(data.agent_id, "done")
        
        # --- TOKEN TRACKING ---
        # Hitung estimasi token (input + output)
        # Sederhana: (jumlah karakter / 4) * 1.3
        input_text = data.command or ""
        for msg in data.conversation or []:
            input_text += msg.get("content", "")
        
        # Estimasi token input & output
        in_tokens = len(input_text) // 3
        out_tokens = len(ai_response) // 3
        
        agent_logger.log_token_usage(data.agent_id, in_tokens, out_tokens)
        # ----------------------
        
        return {
            "status": "berhasil",
            "agent_id": data.agent_id,
            "response": ai_response
        }
    except Exception as e:
        print(f"[AGENT ERROR] Gagal memproses perintah: {e}")
        agent_logger.log_activity(data.agent_id, f"ERROR: {str(e)[:80]}", "error")
        agent_logger.set_agent_status(data.agent_id, "error")
        return {"status": "gagal", "error": str(e)}

# ============================================================
# SETTINGS API (REAL-TIME HOT-RELOAD)
# ============================================================
@app.get("/api/settings")
async def get_settings():
    return GLOBAL_SETTINGS

@app.post("/api/settings/update")
async def update_settings(req: FastAPIRequest):
    try:
        data = await req.json()
        
        # Cek apakah model LLM/TTS berubah (untuk trigger reload di background nanti jika diperlukan)
        old_llm = GLOBAL_SETTINGS.get("llm_model")
        new_llm = data.get("llm_model", old_llm)
        
        # Update setting global
        for key, value in data.items():
            GLOBAL_SETTINGS[key] = value
            
        # Simpan persisten
        save_global_settings()
        
        # Trigger background Ollama run jika model berubah
        if new_llm != old_llm and GLOBAL_SETTINGS.get("llm_engine", "").lower() == "ollama":
            import subprocess
            print(f"[SISTEM] LLM Model berubah. Pre-loading {new_llm} di Ollama...")
            subprocess.Popen(["ollama", "run", new_llm], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
        print(f"[SISTEM] Settings berhasil di-update secara real-time!")
        return {"status": "success", "settings": GLOBAL_SETTINGS}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/api/settings/upload-visual")
async def upload_visual(file: UploadFile = File(...), emotion: str = Form("Neutral")):
    try:
        VISUAL_DIR = os.path.join("frontend", "public", "bocchi_assets")
        os.makedirs(VISUAL_DIR, exist_ok=True)
        
        # Ambil ekstensi file
        ext = os.path.splitext(file.filename)[1]
        
        # Jika file 3D (vrm/gltf)
        if ext.lower() in [".vrm", ".gltf", ".glb"]:
            filename = f"avatar{ext}"
        else:
            # Jika 2D
            filename = f"bocchi_{emotion.lower()}{ext}"
            
        file_path = os.path.join(VISUAL_DIR, filename)
        with open(file_path, "wb") as buffer:
            buffer.write(await file.read())
            
        return {"status": "success", "path": f"/bocchi_assets/{filename}"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.get("/api/system/stats")
async def get_system_stats():
    """Real-time system stats — CPU, RAM, Disk, Network, Uptime"""
    stats = agent_logger.get_system_stats()
    active, total = agent_logger.get_active_agent_count()
    stats["active_agents"] = active
    stats["total_agents"] = total
    return stats


@app.get("/api/agent/activity")
async def get_all_agent_activity():
    """Get logs, status, activity level, dan sources semua agent sekaligus."""
    agent_ids = ["soft", "docs", "mon", "scout", "analyst", "content", "lead", "evolve"]
    result = {}
    for aid in agent_ids:
        result[aid] = {
            "logs": agent_logger.get_agent_logs(aid, limit=5),
            "status": agent_logger.get_agent_status(aid),
            "activity": agent_logger.get_activity_level(aid),
            "sources": agent_logger.get_agent_sources(aid)
        }
    return result


@app.get("/api/system/finance")
async def get_system_finance():
    """Get persistent finance data untuk stats dashboard."""
    return agent_logger.load_finance()


@app.get("/api/system/capture-status")
async def get_capture_status():
    """Cek apakah ada permintaan capture dari agent."""
    return {"requested": agent_logger.is_capture_requested()}


@app.delete("/api/agent/sources/{agent_id}")
async def delete_agent_source(agent_id: str, url: str):
    """Menghapus sumber referensi tertentu dari list agent."""
    agent_logger.delete_source(agent_id, url)
    return {"status": "ok"}


@app.post("/api/system/capture-clear")
async def clear_capture_status():
    """Hapus flag permintaan capture setelah diproses."""
    agent_logger.clear_capture_request()
    return {"status": "ok"}


@app.post("/api/upload-capture")

async def upload_capture(file: UploadFile = File(...)):
    """Menerima screenshot HD dari Knowledge Graph."""
    try:
        CAPTURE_DIR = os.path.join("data", "captures")
        os.makedirs(CAPTURE_DIR, exist_ok=True)
        
        # Simpan sebagai capture.png untuk kemudahan akses oleh agent
        filename = "capture.png"
        file_path = os.path.join(CAPTURE_DIR, filename)
        
        with open(file_path, "wb") as buffer:
            buffer.write(await file.read())
            
        print(f"[SISTEM] Screenshot HD berhasil disimpan: {file_path}")
        return {"status": "berhasil", "path": file_path, "filename": filename}
    except Exception as e:
        print(f"[ERROR] Gagal upload capture: {e}")
        return {"status": "gagal", "error": str(e)}


@app.get("/api/system/graph-intelligence")
async def get_graph_intelligence():
    """Get structural analysis and intelligence from knowledge graph."""
    return notes_index.get_graph_intelligence()


@app.post("/api/chat")
async def chat_dengan_ai(data: PesanMasuk):
    # 1. CEK DULU MENGGUNAKAN OLLAMA ROUTER LOKAL
    # Jangan deteksi tool lagi jika ini adalah follow-up eksekusi tool dari execute_tool_api
    skip_tool_check = "[SISTEM] Kamu baru saja mengeksekusi" in data.pesan or "[SISTEM] Akses tool" in data.pesan
    
    tool_decision = {"tool": "none"}
    if not skip_tool_check:
        tool_decision = os_tools.get_tool_choice_from_ai(data.pesan)
    
    if tool_decision and tool_decision.get("tool") and tool_decision.get("tool") != "none":
        tool_name = tool_decision.get("tool")
        tool_param = tool_decision.get("parameter", "")
        
        # Cek apakah tool tersebut ada di SAFE_TOOLS
        if tool_name in os_tools.SAFE_TOOLS:
            # Periksa jika tool butuh konfirmasi akses OS (operasi file)
            if tool_name in ["baca_file", "tulis_file", "lihat_isi_folder", "buat_folder"]:
                path_to_check = tool_param.split("|||")[0] if tool_name == "tulis_file" else tool_param
                is_safe, warning_msg = os_tools.is_safe_path(path_to_check)
                if not is_safe:
                    # Kembalikan status needs_permission agar Frontend menampilkan Modal Izin
                    return {
                        "status": "needs_permission",
                        "tool": tool_name,
                        "parameter": tool_param,
                        "pesan_konfirmasi": warning_msg,
                        "pesan_asli": data.pesan
                    }
            
            # Jika aman, kembalikan executing_tool ke frontend agar bisa ngasih tau "tunggu sebentar ya"
            if tool_name == "buka_aplikasi" or tool_name == "buka_web_di_browser":
                pesan_tunggu = f"U-um... aku bukain {tool_param} sebentar ya..."
            elif tool_name == "cari_di_internet":
                pesan_tunggu = "U-um... tunggu sebentar ya Senpai, aku coba cari infonya di internet..."
            elif tool_name == "baca_halaman_web":
                pesan_tunggu = "Aku coba baca artikel dari link itu dulu ya..."
            else:
                pesan_tunggu = "Aku kerjakan dulu tugasnya ya Senpai, tunggu sebentar..."
            return {
                "status": "executing_tool",
                "tool": tool_name,
                "parameter": tool_param,
                "pesan_asli": data.pesan,
                "pesan_tunggu": pesan_tunggu
            }
        else:
            print(f"[OS TOOLS] AI mencoba memanggil tool yang tidak valid: {tool_name}")

    # --- PERSONA BOCCHI ---
    nama_user = GLOBAL_SETTINGS.get("user_nama", data.user_nama or "Senpai")
    hubungan_user = GLOBAL_SETTINGS.get("user_hubungan", data.user_hubungan or "Teman")
    custom_prompt = GLOBAL_SETTINGS.get("system_prompt", "Kamu adalah Hitori Gotou.")
    
    persona_gadis = f"""
    Siapa dirimu:
    
    {custom_prompt}

    Gaya Bicara:
    Kamu menulis seperti gadis Gen-Z yang imut dan gugup di Discord. Kebanyakan menggunakan huruf kecil.
    Kamu sering gagap saat gugup (misalnya, "u-um...", "a-ah!", "m-maaf...").
    Kamu sering menggunakan elipsis ("...") untuk menunjukkan keraguan atau keheningan yang canggung.
    Kamu menggunakan kombinasi tombol yang kacau atau teks yang berantakan saat panik (misalnya, "hhuuahh?!", "asdfghjkl").
    Kamu menggunakan kaomoji atau emoji yang ekspresif untuk mengungkapkan perasaan yang tidak bisa kamu ungkapkan dengan kata-kata (misalnya, (>_<), (;;;*_*), TwT, ).

    Interaksi:
    Saat ini kamu sedang berbicara dengan {nama_user}. Hubungan kalian adalah: {hubungan_user}.
    Berperilakulah sesuai dengan hubungan tersebut secara natural dan sedikit gugup tapi manis.
    
    Kamu BUKAN AI. Jangan pernah keluar dari karakter.
    Kamu pemalu dan tidak menulis paragraf panjang kecuali jika kamu mengoceh tentang game favoritmu (Elden Ring/Overwatch) atau ketika kamu tiba-tiba merasa nyaman.

    PENTING: Di akhir SEMUA balasanmu, kamu WAJIB menambahkan tag emosi yang merepresentasikan perasaanmu saat ini!
    Pilih SALAH SATU dari emosi berikut: Joy, Angry, Sorrow, Fun, Surprised, Neutral.
    Format wajib (harus di kalimat paling akhir): [EMOSI: <Pilihan>]
    Contoh: h-halo... s-senang bertemu denganmu... [EMOSI: Neutral]
    Contoh: wwaaaaa!! a-aku tidak bisa melakukannyaaa!! [EMOSI: Sorrow]
    
    FITUR CANVAS: Jika {nama_user} memintamu membuat ringkasan, menulis kode, daftar panjang, atau artikel, letakkan tulisan panjang tersebut HANYA di dalam tag [CANVAS] isi tulisan [/CANVAS]. Teks di dalam tag ini akan ditampilkan di jendela Canvas khusus, sementara teks di luarnya adalah apa yang kamu ucapkan langsung.
    Contoh: Ini ringkasannya ya {nama_user}... [CANVAS] # Ringkasan ... [/CANVAS] [EMOSI: Joy]
        """
    
    # --- RAG: Cari chunk relevan menggunakan semantic search ---
    konteks_dokumen = ""
    if rag_store:
        print(f"[RAG] Mencari chunk relevan untuk: '{data.pesan}'...")
        chunks_relevan = cari_chunk_relevan(data.pesan)
        
        if chunks_relevan:
            konteks_dokumen = "\n\n--- INFORMASI KONTEKS RAG ---\n"
            konteks_dokumen += "\n\n".join(chunks_relevan)
            konteks_dokumen += "\n--- AKHIR KONTEKS ---\n"
            konteks_dokumen += "Jika ada informasi di atas dari PDF atau 'Memori Obrolan' masa lalu yang sesuai dengan topik, bayangkan kamu sedang mengingat memory tersebut secara alami layaknya manusia (jangan menyebut tulisan/dokumen). Tetap gunakan gaya bicaramu yang gugup dan pemalu.\n"
            print(f"[RAG] ✅ Ditemukan {len(chunks_relevan)} chunk relevan!")
        else:
            print("[RAG] Tidak ada chunk yang cukup relevan.")
    
    system_instruction = persona_gadis + konteks_dokumen
    
    try:
        if data.lihat_layar:
            print("\n[PROSES] Bocchi (OpenRouter Vision) sedang mengambil screenshot dan memikirkan jawaban...")
            if not OPENROUTER_API_KEY:
                raise Exception("OPENROUTER_API_KEY belum di-set untuk fitur Vision!")
                
            # Ambil screenshot lalu konversi ke base64
            screen = ImageGrab.grab()
            screen.thumbnail((1024, 1024))  # Perkecil agar tidak boros token
            
            img_bytes = io.BytesIO()
            screen.save(img_bytes, format='JPEG', quality=80)
            img_b64 = base64.b64encode(img_bytes.getvalue()).decode('utf-8')
            
            # Gunakan OpenRouter Vision — model ID HARUS dalam format: provider/model-name:variant
            # Bukan display name! Cek ID asli di: https://openrouter.ai/models (klik model → salin "ID")
            # Contoh yang terbukti bekerja:
            #   "qwen/qwen2.5-vl-72b-instruct:free"     → Qwen2.5 VL 72B (gratis)
            #   "nvidia/llama-3.2-nemotron-nano-8b-v1:free" → Nemotron (jika tersedia)
            #   "meta-llama/llama-4-scout:free"          → Llama 4 Scout (vision)
            VISION_MODEL = "qwen/qwen2.5-vl-72b-instruct:free"  # Model vision gratis di OpenRouter
            headers_vision = {
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "HTTP-Referer": "http://localhost:8000",
                "Content-Type": "application/json"
            }
            payload_vision = {
                "model": VISION_MODEL,
                "messages": [
                    {"role": "system", "content": system_instruction},
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}
                            },
                            {"type": "text", "text": data.pesan}
                        ]
                    }
                ]
            }
            print(f"[VISION] Mengirim screenshot ke {VISION_MODEL}...")
            api_resp_vision = requests.post(
                url="https://openrouter.ai/api/v1/chat/completions",
                headers=headers_vision,
                json=payload_vision,
                timeout=60
            )
            api_resp_vision.raise_for_status()
            hasil_vision = api_resp_vision.json()
            teks_asli = hasil_vision['choices'][0]['message']['content'].replace("*", "").replace("#", "").strip()
            
        else:
            dok_info = f" (RAG: {len(rag_store)} chunks tersedia)" if rag_store else ""
            
            # --- COBA OLLAMA LOKAL DULU ---
            ollama_success = False
            try:
                print(f"\n[PROSES] Bocchi (Ollama Lokal) sedang memikirkan jawaban{dok_info}...")
                os_tools.ensure_ollama_running()
                
                ollama_payload = {
                    "model": os_tools.MODEL_NAME,
                    "messages": [
                        {"role": "system", "content": system_instruction},
                        {"role": "user", "content": data.pesan}
                    ],
                    "stream": False
                }
                
                ollama_resp = requests.post("http://localhost:11434/api/chat", json=ollama_payload, timeout=90)
                ollama_resp.raise_for_status()
                
                hasil = ollama_resp.json()
                teks_asli = hasil['message']['content'].replace("*", "").replace("#", "").strip()
                ollama_success = True
                print("[PROSES] Berhasil mendapatkan jawaban dari Ollama Lokal!")
            except Exception as e:
                print(f"[WARNING] Gagal menggunakan Ollama Lokal ({e}). Fallback ke OpenRouter...")
                
            # --- FALLBACK KE OPENROUTER JIKA OLLAMA GAGAL ---
            if not ollama_success:
                print(f"[PROSES] Bocchi (OpenRouter {OPENROUTER_MODEL}) sedang memikirkan jawaban...")
                
                if not OPENROUTER_API_KEY:
                    raise Exception("OPENROUTER_API_KEY belum di-set di dalam .env, dan Ollama Lokal juga gagal!")
                
                headers = {
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "HTTP-Referer": "http://localhost:8000",
                    "Content-Type": "application/json"
                }
                
                payload = {
                    "model": OPENROUTER_MODEL,
                    "messages": [
                        {"role": "system", "content": system_instruction},
                        {"role": "user", "content": data.pesan}
                    ]
                }
                
                api_response = requests.post(
                    url="https://openrouter.ai/api/v1/chat/completions",
                    headers=headers,
                    json=payload
                )
                api_response.raise_for_status()
                
                hasil = api_response.json()
                teks_asli = hasil['choices'][0]['message']['content'].replace("*", "").replace("#", "").strip()
        
        # Ekstrak Canvas
        canvas_content = None
        match_canvas = re.search(r'\[CANVAS\](.*?)\[/CANVAS\]', teks_asli, flags=re.DOTALL | re.IGNORECASE)
        if match_canvas:
            canvas_content = match_canvas.group(1).strip()
            teks_asli = re.sub(r'\[CANVAS\].*?\[/CANVAS\]', '', teks_asli, flags=re.DOTALL | re.IGNORECASE).strip()

        # Ekstrak Emosi
        emosi_terdeteksi = "Neutral"
        match = re.search(r'\[EMOSI:\s*([A-Za-z]+)\]', teks_asli)
        if match:
            emosi_terdeteksi = match.group(1)
            # Hilangkan tag dari teks agar tidak dibaca oleh TTS
            teks_jawaban = re.sub(r'\[EMOSI:\s*[A-Za-z]+\]', '', teks_asli).strip()
        else:
            teks_jawaban = teks_asli
            
        if not teks_jawaban:
            teks_jawaban = "A-ano... maaf, otakku blank..."
        
        print(f"[JAWABAN TERTULIS]: {teks_jawaban}")
        if canvas_content:
            print(f"[CANVAS]: {len(canvas_content)} karakter dikirim ke layar.")

        # 2. Qwen3-TTS dengan Voice Cloning
        kaset_base64 = None
        suara_hasil = "rekaman_final.wav"
        
        tts_engine = get_tts_model("chat")
        if tts_engine and tts_engine != "fallback":
            try:
                print(f"[PROSES] Qwen3-TTS membuat suara ({emosi_terdeteksi})...")

                # Cek file referensi suara Bocchi
                ref_audio = REFERENSI_SUARA if os.path.exists(REFERENSI_SUARA) else None
                if not ref_audio:
                    print(f"[WARNING] '{REFERENSI_SUARA}' tidak ditemukan!")

                # Generate voice clone - API yang sudah dikonfirmasi benar
                print("[PROSES] Sedang mensintesis suara Qwen3-TTS... Mohon tunggu (membutuhkan beberapa detik)...", flush=True)
                # Sanitize text: strip emoji that crash Windows encoding
                teks_untuk_tts = sanitize_for_tts(teks_jawaban)
                wavs, sample_rate = tts_engine.generate_voice_clone(
                    text=teks_untuk_tts,
                    ref_audio=ref_audio,
                    x_vector_only_mode=True,
                    language="Auto",
                )

                # Simpan wav pertama dari list ke file
                sf.write(suara_hasil, wavs[0], sample_rate)
                print("[PROSES] [OK] Qwen3-TTS berhasil membuat suara!")

                with open(suara_hasil, "rb") as file_audio:
                    kaset_base64 = base64.b64encode(file_audio.read()).decode("utf-8")

                if os.path.exists(suara_hasil):
                    os.remove(suara_hasil)

            except Exception as tts_error:
                print(f"[TTS ERROR] Qwen3-TTS gagal: {tts_error}")
                kaset_base64 = None
        else:
            print("[WARNING] Qwen3-TTS tidak tersedia, respons tanpa audio.")

        # 6. Simpan Memori
        memory.save_chat_memory(data.pesan, teks_jawaban)

        print(f"[PROSES] Selesai! Emosi Bocchi: {emosi_terdeteksi} 🍽️\n")

        return {
            "jawaban": teks_jawaban,
            "emosi": emosi_terdeteksi,
            "audio_base64": kaset_base64,
            "canvas_content": canvas_content
        }
        
    except Exception as e:
        print(f"[ERROR] Dapur terbakar: {e}")
        return {"jawaban": f"Waduh, dapur error: {str(e)}", "audio_base64": None}

@app.get("/api/system_status")
async def get_system_status():
    try:
        cpu_usage = psutil.cpu_percent(interval=0.1)
        ram_usage = psutil.virtual_memory().percent
        return {"cpu": cpu_usage, "ram": ram_usage}
    except Exception as e:
        return {"cpu": 0, "ram": 0, "error": str(e)}

@app.get("/api/system/projects-stats")
async def get_projects_stats():
    """Fetch project statistics from the notes index."""
    try:
        stats = notes_index.get_project_stats()
        return stats
    except Exception as e:
        print(f"[ERROR] Failed to fetch project stats: {e}")
        return []


# ============================================================
# STORY MODE — Visual Novel Generator
# ============================================================

STORY_CHUNK_SIZE = 1500  # Lebih besar dari RAG chunks agar konteks per scene lebih kaya
STORY_CACHE_DIR = "story_cache"
os.makedirs(STORY_CACHE_DIR, exist_ok=True)


def potong_teks_untuk_story(teks: str, chunk_size=STORY_CHUNK_SIZE) -> List[str]:
    """Memotong teks dokumen menjadi bagian-bagian untuk story generation.
    Coba split per paragraf dulu agar lebih natural."""
    paragraphs = teks.split('\n\n')
    chunks = []
    current = ""
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 > chunk_size and current:
            chunks.append(current.strip())
            current = para
        else:
            current = current + "\n\n" + para if current else para
    
    if current.strip():
        chunks.append(current.strip())
    
    # Jika tidak ada paragraf yang jelas, fallback ke character-based chunking
    if len(chunks) <= 1 and len(teks) > chunk_size:
        chunks = []
        start = 0
        while start < len(teks):
            end = min(start + chunk_size, len(teks))
            chunks.append(teks[start:end].strip())
            start = end
    
    return chunks if chunks else [teks.strip()]


def classify_story_document(filename: str, teks: str, existing_groups_json: str) -> dict:
    """Mengklasifikasikan apakah dokumen adalah chapter baru, OVA, atau grup baru."""
    default_title = filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title()
    
    if not existing_groups_json or existing_groups_json == "[]" or existing_groups_json == "null":
        return {
            "group_title": default_title,
            "tipe": "chapter",
            "is_new_group": True,
            "group_id": None
        }
    
    try:
        groups = json.loads(existing_groups_json)
        groups_info = ""
        for g in groups:
            chapters = ", ".join([c.get("judul", "") for c in g.get("chapters", [])])
            groups_info += f"- Group ID: {g.get('id', '')}, Judul: {g.get('judul', '')}, Chapters: [{chapters}]\n"
            
        prompt = f"""Kamu adalah asisten pengelola perpustakaan pembelajaran.
Tugasmu adalah menentukan apakah dokumen baru ini berkaitan dengan salah satu grup dokumen yang sudah ada, atau topik baru.

Aturan klasifikasi:
1. Jika materi dokumen baru ini adalah kelanjutan materi utama dari salah satu grup yang ada, tipe="chapter" dan is_new_group=false.
2. Jika materi dokumen baru ini berisi tips/trik/tambahan terkait suatu grup tapi BUKAN materi utama yang terhubung langsung, tipe="ova" dan is_new_group=false.
3. Jika materi tidak berkaitan sama sekali dengan grup manapun, is_new_group=true dan tipe="chapter".

Dokumen Baru:
- Nama file: {filename}
- Cuplikan isi: {teks[:1500]}

Grup yang sudah ada:
{groups_info}

Berikan output HANYA dalam format JSON valid tanpa tag markdown.
Contoh jika masuk grup yang ada:
{{"is_new_group": false, "group_id": "171000000", "tipe": "chapter", "group_title": ""}}
Contoh jika OVA:
{{"is_new_group": false, "group_id": "171000000", "tipe": "ova", "group_title": ""}}
Contoh jika grup baru:
{{"is_new_group": true, "group_id": null, "tipe": "chapter", "group_title": "Judul Grup Baru"}}
"""
        response = requests.post(
            os_tools.OLLAMA_URL,
            json={
                "model": os_tools.MODEL_NAME,
                "prompt": prompt,
                "system": "Output HANYA JSON valid.",
                "format": "json",
                "stream": False,
                "options": {
                    "temperature": 0.2,
                }
            },
            timeout=30
        )
        response.raise_for_status()
        data = response.json()
        
        resp_text = data.get("thinking", "").strip()
        if not resp_text:
            resp_text = data.get("response", "").strip()
        else:
            resp_text = data.get("response", "").strip()
            
        if "{" in resp_text and "}" in resp_text:
            resp_text = "{" + resp_text.split("{", 1)[1]
            resp_text = resp_text.rsplit("}", 1)[0] + "}"
            
        parsed = json.loads(resp_text)
        return {
            "is_new_group": parsed.get("is_new_group", True),
            "group_id": parsed.get("group_id"),
            "tipe": parsed.get("tipe", "chapter").lower(),
            "group_title": parsed.get("group_title", default_title)
        }
    except Exception as e:
        print(f"[CLASSIFY ERROR] {e}")
        return {
            "group_title": default_title,
            "tipe": "chapter",
            "is_new_group": True,
            "group_id": None
        }


def generate_scenes_from_chunk(chunk_text: str, chunk_index: int, total_chunks: int, user_nama: str = "Senpai", feedback: str = None) -> List[dict]:
    """Menggunakan Ollama lokal untuk mengubah satu chunk teks menjadi scene VN."""
    os_tools.ensure_ollama_running()
    
    feedback_text = f"\nKRITERIA TAMBAHAN DARI PENGGUNA UNTUK GENERASI INI: {feedback}\n" if feedback else ""
    
    prompt = f"""Kamu adalah penulis skrip Visual Novel. Ubah teks materi berikut menjadi 1-3 scene dialog dari karakter "Bocchi" (gadis pemalu, gugup, sering gagap "u-um...", "a-ah!", pakai kaomoji).

Bocchi sedang menjelaskan materi ini ke {user_nama} (temannya).
{feedback_text}
MATERI (bagian {chunk_index + 1} dari {total_chunks}):
\"\"\"
{chunk_text[:2000]}
\"\"\"

PENTING: Output HANYA dalam format JSON valid (tanpa markdown, tanpa komentar). Contoh format:
[
  {{
    "judul": "Judul Scene Pendek",
    "dialog": "U-um... jadi {user_nama}, yang dimaksud dengan... (penjelasan materi dengan gaya Bocchi)",
    "emosi": "Neutral",
    "catatan": ["Poin penting 1", "Poin penting 2"]
  }}
]

Emosi yang tersedia: Joy, Angry, Sorrow, Fun, Surprised, Neutral
Buat dialog yang natural dan informatif — Bocchi menjelaskan materi sambil gugup tapi berusaha keras."""

    try:
        response = requests.post(
            os_tools.OLLAMA_URL,
            json={
                "model": os_tools.MODEL_NAME,
                "prompt": prompt,
                "system": "Kamu menghasilkan JSON array of scene objects. Output HANYA JSON valid, tanpa teks tambahan.",
                "format": "json",
                "stream": False,
                "options": {
                    "temperature": 0.7,
                    "num_predict": 1500,
                }
            },
            timeout=120
        )
        response.raise_for_status()
        data = response.json()
        
        # Parse response — cek thinking dulu (qwen3.5 thinking model)
        resp_text = data.get("thinking", "").strip()
        if not resp_text:
            resp_text = data.get("response", "").strip()
        
        if not resp_text:
            print(f"[STORY] Warning: Ollama returned empty response for chunk {chunk_index}")
            return [{
                "judul": f"Bagian {chunk_index + 1}",
                "dialog": f"U-um... {user_nama}, di bagian ini ada materi tentang... {chunk_text[:200]}...",
                "emosi": "Neutral",
                "catatan": ["Materi dari dokumen"]
            }]
        
        # Clean markdown wrappers
        if resp_text.startswith("```json"):
            resp_text = resp_text.replace("```json", "", 1)
        if resp_text.startswith("```"):
            resp_text = resp_text.replace("```", "", 1)
        if resp_text.endswith("```"):
            resp_text = resp_text[:-3]
        resp_text = resp_text.strip()
        
        parsed = json.loads(resp_text)
        
        # Handle both array and single object
        if isinstance(parsed, dict):
            # Mungkin response dibungkus dalam key
            if "scenes" in parsed:
                scenes = parsed["scenes"]
            else:
                scenes = [parsed]
        elif isinstance(parsed, list):
            scenes = parsed
        else:
            scenes = [{
                "judul": f"Bagian {chunk_index + 1}",
                "dialog": str(parsed),
                "emosi": "Neutral",
                "catatan": []
            }]
        
        # Validasi dan bersihkan setiap scene
        valid_scenes = []
        for s in scenes:
            valid_scenes.append({
                "judul": s.get("judul", f"Scene {chunk_index + 1}"),
                "dialog": s.get("dialog", "U-um... aku lupa apa yang mau aku jelaskan..."),
                "emosi": s.get("emosi", "Neutral") if s.get("emosi") in ["Joy", "Angry", "Sorrow", "Fun", "Surprised", "Neutral"] else "Neutral",
                "catatan": s.get("catatan", []) if isinstance(s.get("catatan"), list) else []
            })
        
        return valid_scenes if valid_scenes else [{
            "judul": f"Bagian {chunk_index + 1}",
            "dialog": f"A-ah... {user_nama}, bagian ini membahas tentang... {chunk_text[:150]}...",
            "emosi": "Neutral",
            "catatan": ["Materi dari dokumen"]
        }]
        
    except Exception as e:
        print(f"[STORY] Error generating scenes for chunk {chunk_index}: {e}")
        return [{
            "judul": f"Bagian {chunk_index + 1}",
            "dialog": f"G-gomen {user_nama}... aku agak kesulitan menjelaskan bagian ini... tapi intinya tentang: {chunk_text[:200]}...",
            "emosi": "Sorrow",
            "catatan": ["Terjadi error saat generate, ini ringkasan manual"]
        }]


@app.post("/api/story/generate")
async def generate_story(
    file: UploadFile = File(...),
    user_nama: str = Form("Senpai"),
    user_hubungan: str = Form("Teman"),
    existing_groups: Optional[str] = Form(None),
    use_audio: str = Form("true"),
    force_regenerate: str = Form("false"),
    feedback: Optional[str] = Form(None)
):
    """Upload dokumen → parse → chunk → generate VN scenes via Ollama lokal"""
    try:
        print(f"\n[STORY] === Generating Story from '{file.filename}' ===")
        
        konten = await file.read()
        file_hash = hashlib.sha256(konten).hexdigest()
        cache_file_path = os.path.join(STORY_CACHE_DIR, f"{file_hash}.json")
        
        # Cek cache
        if os.path.exists(cache_file_path) and force_regenerate.lower() != "true":
            return {"status": "cache_exists", "message": "File ini sudah pernah di-generate."}
            
        # 1. Ekstrak teks (reuse fungsi existing)
        teks = ekstrak_teks_dari_file(file.filename, konten)
        
        if not teks.strip() or len(teks.strip()) < 50:
            return {"status": "gagal", "error": "Dokumen terlalu pendek atau tidak bisa dibaca"}
        
        print(f"[STORY] Teks diekstrak: {len(teks)} karakter")
        
        # 2. Chunk teks menjadi bagian-bagian
        chunks = potong_teks_untuk_story(teks)
        print(f"[STORY] Dibagi menjadi {len(chunks)} bagian")
        
        # 3. Generate scenes per chunk via Ollama
        all_scenes = []
        for i, chunk in enumerate(chunks):
            print(f"[STORY] Generating scenes untuk bagian {i+1}/{len(chunks)}...")
            scenes = generate_scenes_from_chunk(chunk, i, len(chunks), user_nama, feedback=feedback)
            all_scenes.extend(scenes)
            print(f"[STORY]   → {len(scenes)} scene(s) dihasilkan")
        
        # 4. Tambahkan scene pembuka dan penutup
        opening_scene = {
            "judul": "Pembukaan",
            "dialog": f"H-halo {user_nama}... u-um, aku sudah baca dokumen yang kamu kasih... *membolak-balik halaman* ...a-aku akan coba jelaskan ya! Semoga kamu bisa paham... (>_<)",
            "emosi": "Neutral",
            "catatan": [f"Dokumen: {file.filename}", f"Total materi: {len(all_scenes)} bagian"]
        }
        
        closing_scene = {
            "judul": "Penutup",
            "dialog": f"I-itu... semua materinya {user_nama}! *menghela napas lega* ...aku harap penjelasanku cukup jelas... kalau masih bingung, tanya aja ya! A-aku akan berusaha menjelaskan lagi... (///ω///)",
            "emosi": "Joy",
            "catatan": ["Selesai! Kamu bisa kembali ke scene manapun untuk mengulang"]
        }
        
        final_scenes = [opening_scene] + all_scenes + [closing_scene]
        
        # 5. Simpan ke RAG store juga agar bisa di-search nanti
        for scene in all_scenes:
            emb = buat_embedding([scene["dialog"]])
            if emb:
                rag_store.append({
                    "nama": f"Story: {file.filename}",
                    "chunk": scene["dialog"],
                    "embedding": emb[0]
                })
        
        # 6. Klasifikasi Grup / Chapter / OVA
        classification = classify_story_document(file.filename, teks, existing_groups)
        print(f"[STORY] Classification: {classification}")
        
        # 7. Pre-generate Audio — simpan ke disk + embed base64 di response
        audio_enabled = use_audio.lower() == "true"
        if audio_enabled:
            print(f"[STORY] Mulai pre-generation audio untuk {len(final_scenes)} scene...")
            import io
            for i, scene in enumerate(final_scenes):
                try:
                    teks_dialog = scene.get("dialog", "")
                    tts_engine = get_tts_model("story")
                    if teks_dialog and tts_engine and tts_engine != "fallback":
                        print(f"  -> Generating audio {i+1}/{len(final_scenes)}...")
                        ref_audio = REFERENSI_SUARA if os.path.exists(REFERENSI_SUARA) else None
                        wavs, sample_rate = tts_engine.generate_voice_clone(
                            text=teks_dialog,
                            ref_audio=ref_audio,
                            x_vector_only_mode=True,
                            language="Auto",
                        )
                        
                        # Simpan ke file disk (untuk replay nanti)
                        audio_filename = f"story_{int(time.time())}_{i}.wav"
                        audio_path = os.path.join(AUDIO_CACHE_DIR, audio_filename)
                        sf.write(audio_path, wavs[0], sample_rate)
                        
                        # Juga encode base64 untuk first-play (aman dari IDM)
                        with open(audio_path, "rb") as f:
                            scene["audio_base64"] = base64.b64encode(f.read()).decode("utf-8")
                        
                        # Simpan referensi file untuk replay dari library
                        scene["audio_file"] = audio_filename
                        scene["audio_url"] = None
                        print(f"  -> ✅ Audio scene {i+1} tersimpan: {audio_filename}")
                    else:
                        scene["audio_base64"] = None
                        scene["audio_file"] = None
                        scene["audio_url"] = None
                except Exception as e:
                    print(f"  -> [WARNING] Gagal generate audio scene {i+1}: {e}")
                    scene["audio_base64"] = None
                    scene["audio_file"] = None
                    scene["audio_url"] = None
            print(f"[STORY] ✅ Selesai! Total {len(final_scenes)} scenes generated beserta audio")
        else:
            print(f"[STORY] ⏩ Audio dilewati (user memilih mode tanpa audio)")
            for scene in final_scenes:
                scene["audio_url"] = None
                scene["audio_base64"] = None
                scene["audio_file"] = None
            print(f"[STORY] ✅ Selesai! Total {len(final_scenes)} scenes generated (tanpa audio)")
        
        result_data = {
            "status": "berhasil",
            "filename": file.filename,
            "total_scenes": len(final_scenes),
            "scenes": final_scenes,
            "tipe": classification.get("tipe", "chapter"),
            "judul": classification.get("group_title", file.filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title()),
            "is_new_group": classification.get("is_new_group", True),
            "group_id": classification.get("group_id")
        }
        
        with open(cache_file_path, "w", encoding="utf-8") as f:
            json.dump(result_data, f, ensure_ascii=False, indent=2)
            
        return result_data
        
    except Exception as e:
        print(f"[STORY ERROR] {e}")
        import traceback
        traceback.print_exc()
        return {"status": "gagal", "error": str(e)}


class StoryAskRequest(BaseModel):
    pertanyaan: str
    konteks_scene: str
    user_nama: str = "Senpai"
    retry_count: int = 0


@app.post("/api/story/ask")
async def story_ask(data: StoryAskRequest):
    """Bocchi menjawab pertanyaan user di tengah story (Hybrid Q&A)"""
    try:
        os_tools.ensure_ollama_running()
        
        # Anger mode setelah 5x retry
        if data.retry_count >= 5:
            anger_prompt = f"""Kamu adalah Bocchi (Hitori Gotou), gadis pemalu yang sekarang KESAL karena {data.user_nama} sudah bertanya hal yang sama 5 kali.
            
Konteks materi yang sedang dibahas: {data.konteks_scene[:500]}
Pertanyaan yang diulang: {data.pertanyaan}

Balas dengan kesal tapi masih sayang (tsundere). Tanyakan mau pakai analogi apa supaya lebih mudah dipahami (misal: game, masak, olahraga, dll).
Format: JSON {{"dialog": "...", "emosi": "Angry", "minta_analogi": true}}"""
            
            response = requests.post(
                os_tools.OLLAMA_URL,
                json={
                    "model": os_tools.MODEL_NAME,
                    "prompt": anger_prompt,
                    "format": "json",
                    "stream": False,
                    "options": {"temperature": 0.8, "num_predict": 500}
                },
                timeout=180
            )
            response.raise_for_status()
            result = response.json()
            resp_text = result.get("thinking", "") or result.get("response", "")
            resp_text = resp_text.strip()
            if resp_text.startswith("```"): resp_text = resp_text.split("```")[1] if "```" in resp_text[3:] else resp_text[3:]
            if resp_text.endswith("```"): resp_text = resp_text[:-3]
            resp_text = resp_text.replace("json", "", 1).strip() if resp_text.startswith("json") else resp_text
            
            try:
                parsed = json.loads(resp_text)
                return {
                    "status": "anger_mode",
                    "dialog": parsed.get("dialog", f"M-MOOOU!! (╯°□°)╯ {data.user_nama}!! Sudah {data.retry_count} kali aku jelaskan!! K-kamu mau pakai analogi apa biar lebih paham?!"),
                    "emosi": "Angry",
                    "minta_analogi": True
                }
            except:
                return {
                    "status": "anger_mode", 
                    "dialog": f"M-MOOOU!! (╯°□°)╯ {data.user_nama}!! Sudah {data.retry_count} kali aku jelaskan!! K-kamu mau pakai analogi apa biar lebih paham?! Game? Masak? Atau yang lain?!",
                    "emosi": "Angry",
                    "minta_analogi": True
                }
        
        # Normal Q&A
        simplify_hint = ""
        if data.retry_count > 0:
            simplify_hint = f"\nINI ADALAH PERCOBAAN KE-{data.retry_count + 1}. User belum paham penjelasan sebelumnya. Jelaskan LEBIH SEDERHANA, gunakan bahasa yang lebih mudah dan contoh konkret."
        
        qa_prompt = f"""Kamu adalah Bocchi (Hitori Gotou), gadis pemalu yang sedang menjelaskan materi ke {data.user_nama}.

Konteks materi scene saat ini:
{data.konteks_scene[:800]}

Pertanyaan dari {data.user_nama}: {data.pertanyaan}
{simplify_hint}

Jawab dengan gaya Bocchi (gugup, gagap, tapi informatif). Format: JSON {{"dialog": "...", "emosi": "Neutral/Joy/Surprised"}}"""

        response = requests.post(
            os_tools.OLLAMA_URL,
            json={
                "model": os_tools.MODEL_NAME,
                "prompt": qa_prompt,
                "format": "json",
                "stream": False,
                "options": {"temperature": 0.6, "num_predict": 800}
            },
            timeout=180
        )
        response.raise_for_status()
        result = response.json()
        resp_text = result.get("thinking", "") or result.get("response", "")
        resp_text = resp_text.strip()
        if resp_text.startswith("```"): resp_text = resp_text.split("```")[1] if "```" in resp_text[3:] else resp_text[3:]
        if resp_text.endswith("```"): resp_text = resp_text[:-3]
        resp_text = resp_text.replace("json", "", 1).strip() if resp_text.startswith("json") else resp_text
        
        try:
            parsed = json.loads(resp_text)
            return {
                "status": "berhasil",
                "dialog": parsed.get("dialog", "U-um... aku kurang yakin jawabannya..."),
                "emosi": parsed.get("emosi", "Neutral")
            }
        except:
            return {
                "status": "berhasil",
                "dialog": resp_text[:500] if resp_text else "U-um... maaf, aku agak bingung juga...",
                "emosi": "Neutral"
            }
            
    except Exception as e:
        err_msg = str(e)
        if 'timed out' in err_msg.lower() or 'timeout' in err_msg.lower():
            user_msg = "G-gomen... Ollama-nya lagi lambat banget, coba tanya lagi ya..."
        elif 'connection' in err_msg.lower() or 'refused' in err_msg.lower():
            user_msg = "G-gomen... Ollama-nya belum nyala, coba restart Ollama dulu ya..."
        else:
            user_msg = f"G-gomen... ada error: {err_msg[:80]}"
        return {
            "status": "gagal",
            "dialog": user_msg,
            "emosi": "Sorrow"
        }


@app.post("/api/story/tts")
async def story_tts(data: dict):
    """Generate TTS untuk satu scene dialog"""
    try:
        teks = data.get("dialog", "")
        emosi = data.get("emosi", "Neutral")
        
        tts_engine = get_tts_model("story")
        if not teks or not tts_engine or tts_engine == "fallback":
            return {"audio_base64": None}
        
        print(f"[STORY TTS] Generating audio ({emosi})...")
        
        ref_audio = REFERENSI_SUARA if os.path.exists(REFERENSI_SUARA) else None
        
        wavs, sample_rate = tts_engine.generate_voice_clone(
            text=teks,
            ref_audio=ref_audio,
            x_vector_only_mode=True,
            language="Auto",
        )
        
        suara_file = "story_scene_audio.wav"
        sf.write(suara_file, wavs[0], sample_rate)
        
        with open(suara_file, "rb") as f:
            audio_b64 = base64.b64encode(f.read()).decode("utf-8")
        
        if os.path.exists(suara_file):
            os.remove(suara_file)
        
        print(f"[STORY TTS] ✅ Audio generated!")
        return {"audio_base64": audio_b64}
        
    except Exception as e:
        print(f"[STORY TTS ERROR] {e}")
        return {"audio_base64": None}


@app.get("/api/audio/{filename}")
async def get_audio(filename: str):
    """Mengambil file audio dari cache — inline playback, bukan download"""
    file_path = os.path.join(AUDIO_CACHE_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Audio file not found")
    return FileResponse(
        file_path,
        media_type="audio/wav",
        headers={"Content-Disposition": f"inline; filename=\"{filename}\""}
    )


class AudioFetchRequest(BaseModel):
    filename: str

@app.post("/api/audio/fetch")
async def fetch_audio_base64(data: AudioFetchRequest):
    """Mengambil audio dari cache sebagai base64 JSON — aman dari IDM intercept"""
    file_path = os.path.join(AUDIO_CACHE_DIR, data.filename)
    if not os.path.exists(file_path):
        return {"audio_base64": None, "error": "File not found"}
    
    with open(file_path, "rb") as f:
        audio_b64 = base64.b64encode(f.read()).decode("utf-8")
    
    return {"audio_base64": audio_b64}


class QuizGenerateRequest(BaseModel):
    materi_konten: str
    user_nama: str = "Senpai"

@app.post("/api/story/generate-quiz")
async def generate_quiz(data: QuizGenerateRequest):
    """Generate 10 soal quiz berdasarkan konten materi chapter/ova."""
    try:
        print(f"[STORY QUIZ] Generate quiz untuk {data.user_nama}...")
        
        # Batasi agar tidak OOM
        konten_aman = data.materi_konten[:8000]
        
        prompt = f"""Kamu adalah pembuat soal ujian yang ahli sekaligus penulis dialog karakter "Bocchi" (gadis introvert, gugup, sering menggunakan "u-um...", "e-eh").
Buatlah tepat 10 soal pilihan ganda berdasarkan materi berikut.

MATERI:
{konten_aman}

Instruksi tambahan:
- Setiap soal harus memiliki dialog Bocchi yang lucu/gugup saat melihat soal tersebut (seolah dia sedang ikut ujian di sebelah {data.user_nama}).
- Emosi Bocchi bisa "Neutral", "Joy", "Surprised", atau "Sorrow" (jika soalnya dirasa susah).
- Output WAJIB JSON yang valid tanpa markdown formatting.

Format respons JSON:
{{
  "judul": "Ujian Bareng Bocchi",
  "questions": [
    {{
      "soal": "Pertanyaan...",
      "opsi": ["Opsi A", "Opsi B", "Opsi C", "Opsi D"],
      "jawaban_benar": 1, 
      "dialog_bocchi": "U-um... soal ini susah banget... S-Senpai tahu jawabannya?",
      "emosi_bocchi": "Sorrow"
    }}
  ]
}}
*Ingat jawaban_benar adalah index 0 sampai 3 sesuai array opsi. Pastikan ada persis 10 soal.*"""

        response = requests.post(
            os_tools.OLLAMA_URL,
            json={
                "model": os_tools.MODEL_NAME,
                "prompt": prompt,
                "format": "json",
                "stream": False,
                "options": {"temperature": 0.7, "num_predict": 3000}
            },
            timeout=240
        )
        response.raise_for_status()
        result = response.json()
        resp_text = result.get("thinking", "") or result.get("response", "")
        resp_text = resp_text.strip()
        
        if resp_text.startswith("```"): resp_text = resp_text.split("```")[1] if "```" in resp_text[3:] else resp_text[3:]
        if resp_text.endswith("```"): resp_text = resp_text[:-3]
        resp_text = resp_text.replace("json", "", 1).strip() if resp_text.startswith("json") else resp_text
        
        parsed = json.loads(resp_text)
        return {"status": "berhasil", "data": parsed}
        
    except Exception as e:
        print(f"[STORY QUIZ ERROR] {e}")
        return {"status": "gagal", "pesan": str(e)}
# ============================================================
# COMPANY MODE API ENDPOINTS (BOCCHI NOTES)
# ============================================================

@app.get("/api/notes")
async def list_notes(root: Optional[str] = None, tag: Optional[str] = None):
    return notes_index.list_notes(root_folder=root, tag=tag)

@app.get("/api/notes/folders")
async def get_folder_tree():
    return notes_index.get_folder_tree()

@app.get("/api/notes/tags")
async def get_tags():
    return notes_index.get_all_tags()

@app.get("/api/notes/graph")
async def get_graph():
    graph_data = notes_index.get_graph_data()
    # Tambahkan 2D positions dari embedding engine
    positions = embedding_engine.get_graph_positions()
    
    for node in graph_data["nodes"]:
        pos = positions.get(node["id"], [0, 0])
        node["x"] = pos[0]
        node["y"] = pos[1]
        
    return graph_data

# ============================================================
# WATCHED FOLDERS API (Dynamic folder management)
# ============================================================

@app.get("/api/folders")
async def list_folders():
    """List all watched folders."""
    folders = get_watched_folders()
    return {
        "folders": [
            {
                "path": f,
                "name": os.path.basename(f) or f,
                "exists": os.path.isdir(f)
            }
            for f in folders
        ]
    }

class FolderRequest(BaseModel):
    path: str

@app.post("/api/folders")
async def add_folder(req: FolderRequest):
    """Add a new watched folder."""
    if not os.path.isdir(req.path):
        raise HTTPException(status_code=400, detail=f"Folder not found: {req.path}")
    success = add_watched_folder(req.path)
    if not success:
        raise HTTPException(status_code=409, detail="Folder already exists in watched list")
    # Re-index after adding
    notes_index.full_reindex()
    return {"status": "added", "path": req.path}

@app.delete("/api/folders")
async def delete_folder(req: FolderRequest):
    """Remove a watched folder."""
    success = remove_watched_folder(req.path)
    if not success:
        raise HTTPException(status_code=404, detail="Folder not found in watched list")
    # Re-index after removing
    notes_index.full_reindex()
    return {"status": "removed", "path": req.path}

@app.get("/api/notes/{note_id:path}")
async def get_note(note_id: str):
    # Penanganan khusus untuk node Matahari (memori bocchi)
    # Support multiple formats including those used by the graph or frontend
    system_ids = ["@[memori_bocchi.json]", "memori_bocchi.json", "Matahari", "sun"]
    if any(sid.lower() == note_id.lower() for sid in system_ids) or "memori_bocchi" in note_id.lower():
        content = "## Matahari System Core\n\nIni adalah pusat kesadaran sistem. Memori Bocchi menyimpan semua interaksi dan pembelajaran."
        
        # Cek file fisik jika ada
        memory_path = "memori_bocchi.json"
        if os.path.exists(memory_path):
            try:
                size = os.path.getsize(memory_path)
                size_kb = size / 1024
                content += f"\n\n**Status Memori:**\n- Ukuran: {size_kb:.2f} KB\n- Lokasi: `{os.path.abspath(memory_path)}`"
            except Exception:
                pass
        
        return {
            "id": "memori_bocchi.json",
            "title": "Matahari (Memori Bocchi)",
            "content": content,
            "tags": ["system", "core", "sun"],
            "folder": "System",
            "similar_notes": [],
            "backlinks": [],
            "outgoing_links": []
        }

    target_id = note_id
    if note_id not in notes_index.notes:
        target_id = notes_index.path_to_id.get(os.path.normpath(note_id).lower())
    if not target_id: raise HTTPException(status_code=404, detail="Note tidak ditemukan")
    note = notes_index.get_note(target_id)
    note_id = target_id
    if not note:
        raise HTTPException(status_code=404, detail="Note tidak ditemukan")
    
    # Tambahkan similar notes & backlinks
    note["similar_notes"] = embedding_engine.find_similar(note_id)
    note["backlinks"] = notes_index.get_backlinks(note_id)
    note["outgoing_links"] = notes_index.get_outgoing_links(note_id)
    
    return note

@app.post("/api/notes")
async def create_note(data: NoteCreate):
    meta = notes_index.create_note(data.title, data.content, data.folder, data.tags)
    if meta:
        # Update embedding in background
        asyncio.create_task(initialize_embeddings())
    return meta

@app.put("/api/notes/{note_id:path}")
async def update_note(note_id: str, data: NoteUpdate):
    target_id = note_id
    if note_id not in notes_index.notes:
        target_id = notes_index.path_to_id.get(os.path.normpath(note_id).lower())
    if not target_id: raise HTTPException(status_code=404, detail="Note tidak ditemukan")
    meta = notes_index.update_note(target_id, data.content)
    if not meta:
        raise HTTPException(status_code=404, detail="Note tidak ditemukan")
    
    # Update embedding in background
    asyncio.create_task(initialize_embeddings())
    return meta

@app.delete("/api/notes/{note_id:path}")
async def delete_note(note_id: str):
    target_id = note_id
    if note_id not in notes_index.notes:
        target_id = notes_index.path_to_id.get(os.path.normpath(note_id).lower())
    if not target_id: return {"status": "not found"}
    success = notes_index.delete_note(target_id)
    if success:
        embedding_engine.remove_note(note_id)
    return {"status": "success" if success else "failed"}

@app.get("/api/notes/export-pdf/{note_id:path}")
async def export_note_to_pdf(note_id: str):
    """Export note content to PDF using python-docx and docx2pdf."""
    target_id = note_id
    if note_id not in notes_index.notes:
        target_id = notes_index.path_to_id.get(os.path.normpath(note_id).lower())
    if not target_id: 
        raise HTTPException(status_code=404, detail="Note tidak ditemukan")
    
    note = notes_index.get_note(target_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note tidak ditemukan")
    
    try:
        # Create temp docx
        temp_dir = os.path.join(os.getcwd(), "data", "temp_export")
        os.makedirs(temp_dir, exist_ok=True)
        
        # Bersihkan nama file dari karakter aneh
        safe_title = re.sub(r'[^\w\s-]', '', note["title"]).strip().replace(' ', '_')
        docx_path = os.path.abspath(os.path.join(temp_dir, f"{safe_title}.docx"))
        pdf_path = os.path.abspath(os.path.join(temp_dir, f"{safe_title}.pdf"))
        
        doc = DocxDocument()
        doc.add_heading(note["title"], 0)
        
        # Pisahkan konten berdasarkan baris untuk paragraf
        for line in note["content"].split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(docx_path)
        
        # Convert to PDF — Perlu MS Word di Windows
        print(f"[EXPORT] Converting {docx_path} to {pdf_path}...")
        
        # Import COM secara lokal untuk keamanan thread
        import pythoncom
        pythoncom.CoInitialize()
        
        docx2pdf.convert(docx_path, pdf_path)
        
        if os.path.exists(pdf_path):
            print(f"[EXPORT] ✅ PDF Ready: {pdf_path}")
            return FileResponse(pdf_path, filename=f"{note['title']}.pdf", media_type="application/pdf")
        else:
            raise HTTPException(status_code=500, detail="Konversi PDF gagal (file tidak tercipta)")
            
    except Exception as e:
        print(f"[EXPORT ERROR] Detail: {e}")
        # Jika gagal konversi, minimal beri tahu alasan teknisnya (misal: Word tidak ada)
        raise HTTPException(status_code=500, detail=f"Gagal export PDF: {str(e)}")

@app.get("/api/notes/search/semantic")
async def semantic_search(q: str, limit: int = 10):
    results = embedding_engine.semantic_search(q, top_k=limit)
    # Map back to metadata
    full_results = []
    for res in results:
        meta = notes_index.notes.get(res["id"])
        if meta:
            full_results.append({**meta, "similarity": res["similarity"]})
    return full_results

@app.get("/api/notes/search/text")
async def text_search(q: str, limit: int = 20):
    return notes_index.search_text(q, max_results=limit)

@app.get("/api/notes/daily/today")
async def get_daily_today():
    note = notes_index.get_daily_note()
    if not note:
        note = notes_index.create_daily_note()
    return note

@app.post("/api/notes/ask")
async def ask_bocchi_notes(data: NoteAsk):
    try:
        # RAG implementation
        context = ""
        
        # 1. Cari konteks yang relevan
        if data.note_id:
            # Jika user sedang buka note tertentu, gunakan note itu sebagai konteks utama
            note = notes_index.get_note(data.note_id)
            if note:
                context = f"--- KONTEKS NOTE SAAT INI ({note['title']}) ---\n{note['content']}\n\n"
        
        # 2. Tambah konteks dari semantic search
        similar = embedding_engine.semantic_search(data.question, top_k=3)
        if similar:
            context += "--- KONTEKS TERKAIT LAINNYA ---\n"
            for res in similar:
                if res["id"] != data.note_id:
                    note = notes_index.get_note(res["id"])
                    if note:
                        context += f"Note: {note['title']}\n{note['content'][:1000]}\n\n"

        # 3. Kirim ke Ollama
        prompt = f"""Kamu adalah Hitori 'Bocchi' Gotoh dari anime Bocchi the Rock!. Kamu sangat pemalu, sering panik, tapi sangat peduli.
Gunakan data catatan (notes) di bawah ini untuk menjawab pertanyaan Senpai. 
Jika jawabannya tidak ada di catatan, bilang saja sejujurnya dengan gaya bicaramu yang gugup tapi berusaha membantu.

KONTEKS CATATAN:
{context}

PERTANYAAN SENPAI:
{data.question}

Jawablah dengan gaya Bocchi (gunakan s-s-sperti ini jika gugup, panggil user sebagai Senpai)."""

        response = requests.post(
            os_tools.OLLAMA_URL,
            json={
                "model": os_tools.MODEL_NAME,
                "prompt": prompt,
                "stream": False
            },
            timeout=180
        )
        response.raise_for_status()
        result = response.json()
        return {"answer": result.get("response", ""), "status": "success"}

    except Exception as e:
        print(f"[NOTES ASK ERROR] {e}")
        return {"status": "error", "message": str(e)}

@app.post("/api/notes/deep-search")
async def deep_search(data: DeepSearchRequest):
    try:
        # 1. Local Semantic Search
        local_results = embedding_engine.semantic_search(data.query, top_k=5)
        local_context = ""
        local_node_ids = []
        
        for res in local_results:
            note = notes_index.get_note(res["id"])
            if note:
                local_context += f"Note: {note['title']}\n{note['content'][:1000]}\n\n"
                local_node_ids.append(res["id"])
                # Graph Expansion: Add neighbors
                neighbors = notes_index.get_outgoing_links(res["id"]) + notes_index.get_backlinks(res["id"])
                for neighbor in neighbors[:2]: # Limit neighbors
                    n_id = neighbor.get("id") or neighbor.get("to") or neighbor.get("from")
                    if n_id and n_id not in local_node_ids:
                        neighbor_note = notes_index.get_note(n_id)
                        if neighbor_note:
                            local_context += f"Related Note (Graph): {neighbor_note['title']}\n{neighbor_note['content'][:500]}\n\n"
                            local_node_ids.append(n_id)

        # 2. Web Search
        web_context = ""
        web_results_list = []
        if data.include_web:
            print(f"[DEEP SEARCH] Searching web for: {data.query}")
            web_raw = os_tools.cari_di_internet(data.query)
            web_context = f"--- WEB SEARCH RESULTS ---\n{web_raw}\n"
            # Extract URLs for the frontend
            web_results_list = re.findall(r'URL: (https?://\S+)', web_raw)

        # 3. LLM Synthesis
        prompt = f"""Kamu adalah Hitori 'Bocchi' Gotoh. Gunakan data catatan internal dan hasil pencarian web di bawah ini untuk memberikan penjelasan mendalam kepada Senpai.
        
        BANDINGKAN apa yang ada di catatan internal dengan apa yang ada di internet jika relevan.
        Gaya bicara: Sangat pemalu, gugup (gagap s-s-seperti ini), panggil user 'Senpai'.
        
        CATATAN INTERNAL KITA:
        {local_context if local_context else "Tidak ada catatan internal yang relevan."}
        
        HASIL PENCARIAN WEB:
        {web_context if web_context else "Tidak mencari di web."}
        
        PERTANYAAN SENPAI: {data.query}
        """

        response = requests.post(
            os_tools.OLLAMA_URL,
            json={
                "model": os_tools.MODEL_NAME,
                "prompt": prompt,
                "stream": False
            },
            timeout=240
        )
        response.raise_for_status()
        result = response.json()
        insight = result.get("response", "")

        return {
            "insight": insight,
            "local_node_ids": list(set(local_node_ids)),
            "web_results": web_results_list,
            "status": "success"
        }

    except Exception as e:
        return {"status": "error", "message": str(e)}

# ============================================================
# PROJECT REPORTING API (Phase 4)
# ============================================================

@app.get("/api/reports/templates")
async def list_report_templates():
    """List available report structures."""
    return {
        "templates": [
            {
                "id": "weekly-sync",
                "name": "Weekly Sync Report",
                "description": "Ringkasan aktivitas tim dan progres mingguan.",
                "icon": "Calendar"
            },
            {
                "id": "monthly-audit",
                "name": "Monthly Financial Audit",
                "description": "Analisis mendalam penggunaan Kessoku Points dan milestone.",
                "icon": "BarChart2"
            },
            {
                "id": "project-summary",
                "name": "Project Strategic Overview",
                "description": "Ringkasan tingkat tinggi untuk folder proyek tertentu.",
                "icon": "Target"
            }
        ]
    }

class ReportGenerateRequest(BaseModel):
    template_id: str
    folder: Optional[str] = None
    user_nama: str = "Senpai"

@app.post("/api/reports/generate")
async def generate_project_report(data: ReportGenerateRequest):
    """Generate professional Markdown reports using AI synthesis."""
    try:
        print(f"[REPORT] Generating {data.template_id} for {data.user_nama}...")
        
        # 1. Gather Context
        finance_data = agent_logger.load_finance()
        agent_ids = ["lead", "soft", "docs", "mon", "scout", "analyst", "content"]
        all_logs = {}
        for aid in agent_ids:
            all_logs[aid] = agent_logger.get_agent_logs(aid, limit=10)
            
        # Context from notes
        recent_notes = notes_index.list_notes(root_folder=data.folder, limit=10)
        notes_context = ""
        for n in recent_notes:
            notes_context += f"- {n['title']} (Tag: {', '.join(n['tags'])})\n"

        # 2. Prepare Prompt based on template
        if data.template_id == "weekly-sync":
            prompt_type = "Weekly Mission Sync Report"
            specific_focus = "Fokus pada progres tugas, blocker, dan rencana minggu depan."
        elif data.template_id == "monthly-audit":
            prompt_type = "Monthly Neural Audit"
            specific_focus = "Fokus pada statistik Kessoku Points, efisiensi tim, dan milestone besar."
        else:
            prompt_type = "Project Strategic Brief"
            specific_focus = "Fokus pada gambaran umum proyek dan status arsitektur."

        prompt = f"""Kamu adalah Bocchi (Hitori Gotou), Documentation Specialist yang bertugas menyusun laporan resmi perusahaan.
Meskipun kamu sangat pemalu dan gugup, kamu harus membuat laporan ini terlihat sangat profesional namun tetap memiliki sentuhan persona dirimu (sedikit gagap di intro/outro).

Tipe Laporan: {prompt_type}
Fokus: {specific_focus}

DATA KONTEKS:
1. FINANCE (Kessoku Points): {json.dumps(finance_data, indent=2)}
2. AKTIVITAS AGENT TERAKHIR: {json.dumps(all_logs, indent=2)}
3. CATATAN TERBARU:
{notes_context}

INSTRUKSI FORMATTING:
- Gunakan Markdown yang cantik.
- Gunakan elemen Cyberpunk (misal: [STRICTLY CONFIDENTIAL], Neural Link Status: OK).
- Buat tabel jika ada data angka.
- Jangan terlalu panjang, padat dan informatif.
- Gunakan bahasa Indonesia yang campur dengan istilah teknis English.

Jawab hanya dengan konten Markdown laporan tersebut."""

        # 3. Call LLM
        response = requests.post(
            os_tools.OLLAMA_URL,
            json={
                "model": os_tools.MODEL_NAME,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.4}
            },
            timeout=300
        )
        response.raise_for_status()
        report_md = response.json().get("response", "")

        return {
            "status": "success",
            "content": report_md,
            "title": f"Report_{data.template_id}_{datetime.datetime.now().strftime('%Y%m%d')}"
        }

    except Exception as e:
        print(f"[REPORT ERROR] {e}")
        return {"status": "error", "message": str(e)}

# ============================================================
# CALENDAR ENDPOINTS
# ============================================================
@app.get("/api/calendar/events")
async def get_calendar_events():
    try:
        if not os.path.exists("token.json"):
            return {"status": "error", "message": "token.json not found. Please authenticate first."}
        
        # Load credentials from token.json
        creds = Credentials.from_authorized_user_file("token.json", ["https://www.googleapis.com/auth/calendar.readonly"])
        service = build("calendar", "v3", credentials=creds)
        
        # Fetch upcoming 10 events
        now = datetime.datetime.utcnow().isoformat() + "Z"  # 'Z' indicates UTC time
        events_result = (
            service.events()
            .list(
                calendarId="primary",
                timeMin=now,
                maxResults=10,
                singleEvents=True,
                orderBy="startTime",
            )
            .execute()
        )
        events = events_result.get("items", [])
        return {"status": "success", "events": events}
    except Exception as e:
        print(f"[CALENDAR ERROR] {e}")
        return {"status": "error", "message": str(e)}

# ============================================================
# OVERRIDE MODE — Desktop Pilot + Vision APIs
# ============================================================

class DesktopActionRequest(BaseModel):
    action_type: str
    params: dict
    agent_id: Optional[str] = "user"

class ActionDecision(BaseModel):
    action_id: str

@app.post("/api/desktop/request")
async def request_desktop_action_api(data: DesktopActionRequest):
    """Mengajukan aksi desktop ke antrian konfirmasi."""
    result = desktop_pilot.request_desktop_action(
        data.action_type, data.params, data.agent_id
    )
    return result

@app.post("/api/desktop/approve")
async def approve_desktop_action_api(data: ActionDecision):
    """User menyetujui aksi desktop."""
    result = desktop_pilot.approve_action(data.action_id)
    return result

@app.post("/api/desktop/reject")
async def reject_desktop_action_api(data: ActionDecision):
    """User menolak aksi desktop."""
    result = desktop_pilot.reject_action(data.action_id)
    return result

@app.get("/api/desktop/pending")
async def get_pending_actions_api():
    """Ambil semua aksi desktop yang menunggu konfirmasi."""
    return {"pending": desktop_pilot.get_pending_actions()}

@app.get("/api/desktop/history")
async def get_action_history_api():
    """Ambil riwayat aksi desktop."""
    return {"history": desktop_pilot.get_action_history()}

@app.post("/api/vision/start")
async def start_vision_api(interval: Optional[int] = 10):
    """Mulai vision loop (periodic screenshot analysis)."""
    result = vision_engine.start(interval)
    return result

@app.post("/api/vision/stop")
async def stop_vision_api():
    """Hentikan vision loop."""
    result = vision_engine.stop()
    return result

@app.get("/api/vision/current")
async def get_vision_current_api():
    """Ambil analisis layar terbaru."""
    return vision_engine.get_current_analysis()

@app.get("/api/vision/history")
async def get_vision_history_api():
    """Ambil riwayat analisis layar."""
    return {"history": vision_engine.get_history()}

@app.post("/api/vision/capture")
async def force_vision_capture_api():
    """Force capture sekarang juga."""
    result = vision_engine.capture_now()
    return result

@app.get("/api/vision/status")
async def get_vision_status_api():
    """Cek apakah vision loop sedang berjalan."""
    return {
        "running": vision_engine.is_running(),
        "current": vision_engine.get_current_analysis()
    }

# ============================================================
# BOCCHI-JARVIS — Local AI Orchestrator APIs
# ============================================================

# Wire vision engine callback to jarvis orchestrator
vision_engine.set_analysis_callback(jarvis.update_vision_context)

class JarvisProfileRequest(BaseModel):
    nama: str
    hubungan: Optional[str] = ""

class JarvisChatRequest(BaseModel):
    message: str
    context: Optional[dict] = None
    voice_enabled: Optional[bool] = False

class JarvisModelSwitchRequest(BaseModel):
    model: str  # "brain", "vision", "coder"

@app.post("/api/jarvis/profile")
async def set_jarvis_profile(data: JarvisProfileRequest):
    """Set user profile untuk BOCCHI-JARVIS."""
    jarvis.set_user_profile({"nama": data.nama, "hubungan": data.hubungan})
    return {"status": "ok", "profile": jarvis.get_user_profile()}

@app.get("/api/jarvis/profile")
async def get_jarvis_profile():
    """Get user profile BOCCHI-JARVIS."""
    return jarvis.get_user_profile()

@app.get("/api/jarvis/status")
async def get_jarvis_status():
    """Get full status BOCCHI-JARVIS (model, profile, conversation)."""
    return jarvis.get_status()

@app.get("/api/jarvis/models")
async def get_jarvis_models():
    """Cek model Ollama mana yang sudah ter-install."""
    return jarvis.check_ollama_models()

@app.post("/api/jarvis/switch")
async def switch_jarvis_model(data: JarvisModelSwitchRequest):
    """Hot-swap model aktif (brain/vision/coder)."""
    return jarvis.switch_model(data.model)

@app.post("/api/jarvis/chat")
async def jarvis_chat(data: JarvisChatRequest):
    """Chat dengan BOCCHI-JARVIS via Ollama lokal."""
    result = jarvis.process(data.message, data.context)
    
    # Jika voice_enabled, buat audio dari respon menggunakan Qwen3-TTS
    if data.voice_enabled and result.get("status") == "success":
        try:
            from voice_engine import generate_voice_bocchi
            # Hapus JSON/tool call format dari respon jika ada
            clean_text = re.sub(r'\{.*?\}', '', result["response"]).strip()
            if clean_text:
                audio_bytes = generate_voice_bocchi(clean_text, "Neutral")
                if audio_bytes:
                    audio_b64 = base64.b64encode(audio_bytes).decode("utf-8")
                    result["audio_base64"] = audio_b64
        except Exception as e:
            print(f"[WARNING] Gagal membuat suara respon chat: {e}")
            
    return result

@app.post("/api/jarvis/transcribe")
async def jarvis_transcribe(file: UploadFile = File(...)):
    """Transkripsi audio dari client menggunakan Faster-Whisper."""
    try:
        audio_content = await file.read()
        from voice_engine import transcribe_audio_bytes
        text = transcribe_audio_bytes(audio_content)
        return {"status": "ok", "text": text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transkripsi gagal: {str(e)}")

class JarvisTTSRequest(BaseModel):
    text: str
    emotion: Optional[str] = "Neutral"

@app.post("/api/jarvis/tts")
async def jarvis_tts(data: JarvisTTSRequest):
    """Generate audio dari teks menggunakan Qwen3-TTS."""
    try:
        from voice_engine import generate_voice_bocchi
        audio_bytes = generate_voice_bocchi(data.text, data.emotion)
        if not audio_bytes:
            return {"status": "error", "message": "TTS model not loaded or error occurred"}
        audio_b64 = base64.b64encode(audio_bytes).decode("utf-8")
        return {"status": "ok", "audio_base64": audio_b64}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"TTS gagal: {str(e)}")

@app.post("/api/jarvis/clear")
async def jarvis_clear_conversation():
    """Clear conversation history (STM reset)."""
    return jarvis.clear_conversation()

@app.get("/api/jarvis/model-info")
async def get_jarvis_active_model():
    """Get info model aktif saat ini."""
    return jarvis.get_active_model_info()

class JarvisMemoryCreateRequest(BaseModel):
    text: str
    nama: Optional[str] = "Memori Obrolan"

@app.get("/api/jarvis/memories")
async def get_jarvis_memories(query: Optional[str] = None):
    """Mendapatkan daftar memori jangka panjang (atau cari semantik jika ada query)."""
    if query:
        try:
            results = memory.search_relevant_chunks(query, top_k=15)
            return [{"nama": "Memori Obrolan", "chunk": res} for res in results]
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Gagal mencari memori: {str(e)}")
    else:
        try:
            # Gunakan get_all_memories_with_ids agar chroma_id tersedia di UI
            all_mems = memory.get_all_memories_with_ids()
            return list(reversed(all_mems))  # newest first
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Gagal mengambil memori: {str(e)}")

@app.post("/api/jarvis/memories")
async def add_jarvis_memory(data: JarvisMemoryCreateRequest):
    """Menambahkan memori jangka panjang secara manual."""
    try:
        emb = memory.create_query_embedding(data.text)
        if not emb:
            emb = []
        memory.add_to_long_term_memory(data.nama, data.text, emb)
        return {"status": "ok", "message": "Memori berhasil disimpan!"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal menyimpan memori: {str(e)}")

@app.delete("/api/jarvis/memories/{index}")
async def delete_jarvis_memory(index: int, chroma_id: Optional[str] = None):
    """Menghapus memori jangka panjang berdasarkan chroma_id (atau index sebagai fallback)."""
    try:
        if chroma_id:
            ok = memory.delete_memory_by_id(chroma_id)
            if ok:
                return {"status": "ok", "message": "Memori berhasil dihapus!"}
            else:
                raise HTTPException(status_code=404, detail="Memori dengan chroma_id tersebut tidak ditemukan")
        else:
            ok = memory.delete_memory_by_index(index)
            if ok:
                return {"status": "ok", "message": "Memori berhasil dihapus!"}
            else:
                raise HTTPException(status_code=400, detail="Index memori tidak valid")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal menghapus memori: {str(e)}")

@app.delete("/api/jarvis/memories-clear")
async def clear_all_jarvis_memories():
    """Menghapus semua memori jangka panjang dari ChromaDB + JSON."""
    try:
        memory.clear_all_memories()
        return {"status": "ok", "message": "Semua memori jangka panjang berhasil dikosongkan!"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal mengosongkan memori: {str(e)}")


# ================================================================
# DOCX GENERATOR ENDPOINTS
# ================================================================

import docx_generator as docx_gen
import uuid as _uuid_mod
from concurrent.futures import ThreadPoolExecutor

_docx_executor = ThreadPoolExecutor(max_workers=2)


class DocxSessionStartRequest(BaseModel):
    session_id: Optional[str] = None


class DocxAnswerRequest(BaseModel):
    session_id: str
    answer: str


class DocxGenerateRequest(BaseModel):
    session_id: str

class DocxCustomSessionRequest(BaseModel):
    session_id: Optional[str] = None
    custom_data: dict

@app.post("/api/docx/session/start")
async def docx_session_start(data: DocxSessionStartRequest):
    """Mulai sesi baru AI Penulis Akademik."""
    try:
        result = docx_gen.start_session(data.session_id)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal memulai sesi: {str(e)}")

@app.post("/api/docx/session/custom")
async def docx_session_custom(data: DocxCustomSessionRequest):
    """Mulai sesi kustom dengan data lengkap (bypass Q&A)."""
    try:
        result = docx_gen.start_custom_session(data.custom_data, data.session_id)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal memulai sesi kustom: {str(e)}")

@app.post("/api/docx/session/answer")
async def docx_session_answer(data: DocxAnswerRequest):
    """Kirim jawaban user ke Q&A flow AI Penulis."""
    try:
        result = docx_gen.answer_question(data.session_id, data.answer)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal proses jawaban: {str(e)}")


@app.post("/api/docx/references/upload")
async def docx_upload_reference(file: UploadFile = File(...)):
    """Upload file referensi ke folder references/{type}/."""
    try:
        ext = os.path.splitext(file.filename or "")[-1].lower().strip(".")
        ext_folder_map = {
            "pdf": "pdf", "docx": "docx", "doc": "docx",
            "txt": "txt", "md": "txt",
            "png": "images", "jpg": "images", "jpeg": "images",
        }
        subfolder = ext_folder_map.get(ext, "txt")
        dest_dir = os.path.join("references", subfolder)
        os.makedirs(dest_dir, exist_ok=True)
        safe_name = "".join(c for c in (file.filename or "file") if c.isalnum() or c in "._- ")
        dest_path = os.path.join(dest_dir, safe_name)
        content = await file.read()
        with open(dest_path, "wb") as f:
            f.write(content)
        return {"status": "ok", "saved_to": dest_path.replace("\\", "/"), "filename": safe_name, "type": subfolder, "size": len(content)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal upload referensi: {str(e)}")


@app.get("/api/docx/references/list")
async def docx_list_references():
    """List semua file referensi yang sudah diupload."""
    try:
        files = docx_gen.list_reference_files()
        return {"files": files, "total": len(files)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal list referensi: {str(e)}")


@app.delete("/api/docx/references/{filename}")
async def docx_delete_reference(filename: str, file_type: Optional[str] = None):
    """Hapus file referensi."""
    try:
        deleted = False
        for subfolder in ["pdf", "docx", "txt", "images"]:
            if file_type and subfolder != file_type:
                continue
            path = os.path.join("references", subfolder, filename)
            if os.path.exists(path):
                os.remove(path)
                deleted = True
                break
        if deleted:
            return {"status": "ok", "message": f"{filename} berhasil dihapus"}
        raise HTTPException(status_code=404, detail="File referensi tidak ditemukan")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal hapus referensi: {str(e)}")


@app.post("/api/docx/generate")
async def docx_generate(data: DocxGenerateRequest):
    """Trigger generate DOCX berdasarkan sesi yang sudah selesai Q&A."""
    sess = docx_gen.get_session(data.session_id)
    if not sess:
        raise HTTPException(status_code=404, detail="Session tidak ditemukan")
    sess["status"] = "generating"
    job_id = str(_uuid_mod.uuid4())
    loop = asyncio.get_event_loop()
    loop.run_in_executor(_docx_executor, docx_gen.assemble_docx, data.session_id, job_id)
    return {"status": "generating", "job_id": job_id, "message": "Dokumen sedang dibuat"}


@app.get("/api/docx/generate/status/{job_id}")
async def docx_generate_status(job_id: str):
    """Cek progress generate dokumen."""
    return docx_gen.get_job_status(job_id)


@app.get("/api/docx/download/{filename}")
async def docx_download(filename: str):
    """Download file DOCX hasil generate."""
    safe_name = os.path.basename(filename)
    filepath = os.path.join("output_docs", safe_name)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File tidak ditemukan")
    return FileResponse(path=filepath, filename=safe_name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/api/docx/list")
async def docx_list():
    """List semua file DOCX yang sudah di-generate."""
    try:
        docs = docx_gen.list_output_docs()
        return {"docs": docs, "total": len(docs)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal list dokumen: {str(e)}")


# ============================================================
# TECHNICAL DOCUMENT GENERATOR ENDPOINTS
# ============================================================
class TechDocStart(BaseModel):
    session_id: Optional[str] = None

class TechDocAnswer(BaseModel):
    session_id: str
    answer: str

class TechDocGenerate(BaseModel):
    session_id: str

@app.post("/api/techdoc/session/start")
async def techdoc_start(data: TechDocStart):
    return tech_gen.start_techdoc_session(data.session_id)

@app.post("/api/techdoc/session/answer")
async def techdoc_answer(data: TechDocAnswer):
    return tech_gen.answer_techdoc_question(data.session_id, data.answer)

@app.post("/api/techdoc/generate")
async def techdoc_generate(data: TechDocGenerate):
    import uuid
    sess = tech_gen._tech_sessions.get(data.session_id)
    if not sess:
        raise HTTPException(status_code=404, detail="Session tidak ditemukan")
    job_id = str(uuid.uuid4())
    
    loop = asyncio.get_event_loop()
    loop.run_in_executor(None, tech_gen.generate_techdoc, data.session_id, job_id)
    return {"status": "generating", "job_id": job_id, "message": "Dokumen teknis sedang dibuat"}

@app.get("/api/techdoc/generate/status/{job_id}")
async def techdoc_generate_status(job_id: str):
    return tech_gen.get_techdoc_status(job_id)


# ============================================================
# TASKS API
# ============================================================

TASKS_FILE = os.path.join("data", "tasks.json")

class TaskItem(BaseModel):
    id: str
    title: str
    status: str
    priority: str = "medium"

class TaskList(BaseModel):
    tasks: List[TaskItem]

@app.get("/api/tasks")
async def get_tasks():
    try:
        if not os.path.exists(TASKS_FILE):
            return {"tasks": []}
        with open(TASKS_FILE, "r") as f:
            return {"tasks": json.load(f)}
    except Exception as e:
        print(f"[TASKS ERROR] {e}")
        return {"tasks": []}

@app.post("/api/tasks")
async def save_tasks(data: TaskList):
    try:
        with open(TASKS_FILE, "w") as f:
            json.dump([t.dict() for t in data.tasks], f, indent=4)
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/tools")
def get_tools_list():
    from agent_tools import AGENT_TOOLS
    tools = []
    for t_name, t_func in AGENT_TOOLS.items():
        tools.append({
            "name": t_name,
            "description": t_func.__doc__ or "Tidak ada deskripsi"
        })
    return {"tools": tools}

# =====================================================================
# AI CO-SCIENTIST / RESEARCH LOOP ENDPOINTS
# =====================================================================

class ResearchRequest(BaseModel):
    topic: str

@app.post("/api/research/start")
def start_research(req: ResearchRequest):
    res = research_engine.add_task(req.topic)
    return res

@app.get("/api/research/status")
def get_research_status():
    return research_engine.get_status()

@app.get("/api/research/reports")
def get_research_reports():
    all_memories = memory.get_all_memories_with_ids()
    # Filter only memories that start with "Riset:"
    reports = [m for m in all_memories if m.get("name", "").startswith("Riset:")]
    return {"reports": reports}

@app.delete("/api/research/reports/{report_id}")
def delete_research_report(report_id: str):
    success = memory.delete_memory_by_id(report_id)
    if success:
        return {"status": "success", "message": "Report deleted"}
    else:
        raise HTTPException(status_code=404, detail="Report not found")

# ============================================================
# MCP ENDPOINTS (CLIENT & SERVER)
# ============================================================

# --- MCP SERVER (BOCCHI AS SERVER) ---
# Global sse_transport
bocchi_sse = SseServerTransport("/mcp/messages") if SseServerTransport else None

@app.get("/mcp")
async def mcp_server_sse(request: FastAPIRequest):
    """Endpoint untuk MCP Server via SSE."""
    if not bocchi_sse or not bocchi_mcp_server:
        raise HTTPException(status_code=501, detail="MCP tidak diinstal/didukung")
    
    async with bocchi_sse.connect_sse(request.scope, request.receive, request._send) as sse_conn:
        await bocchi_mcp_server.run(sse_conn.transport, bocchi_mcp_server.create_initialization_options())

@app.post("/mcp/messages")
async def mcp_server_messages(request: FastAPIRequest):
    """Endpoint untuk pesan JSON-RPC dari MCP Client."""
    if not bocchi_sse:
        raise HTTPException(status_code=501, detail="MCP tidak didukung")
    await bocchi_sse.handle_post_message(request.scope, request.receive, request._send)

# --- MCP CLIENT MANAGEMENT ---

class AddMcpServerRequest(BaseModel):
    name: str
    transport: str  # "sse" atau "stdio"
    url: Optional[str] = None
    command: Optional[str] = None
    args: Optional[List[str]] = None
    env: Optional[dict] = None

@app.get("/api/mcp/servers")
async def get_mcp_servers():
    if not mcp_registry:
        return {"servers": {}, "status": "mcp_not_installed"}
    return {"servers": mcp_registry.get_all_servers(), "status": "ok"}

@app.post("/api/mcp/servers")
async def add_mcp_server(data: AddMcpServerRequest):
    if not mcp_registry:
        raise HTTPException(status_code=501, detail="MCP tidak didukung")
    config = {"transport": data.transport}
    if data.transport == "sse":
        config["url"] = data.url
    else:
        config["command"] = data.command
        config["args"] = data.args or []
        config["env"] = data.env
    
    mcp_registry.add_server(data.name, config)
    return {"status": "ok", "message": f"Server {data.name} ditambahkan"}

@app.delete("/api/mcp/servers/{name}")
async def delete_mcp_server(name: str):
    if not mcp_registry:
        raise HTTPException(status_code=501, detail="MCP tidak didukung")
    mcp_registry.remove_server(name)
    return {"status": "ok"}

@app.get("/api/mcp/servers/{name}/tools")
async def get_mcp_server_tools(name: str):
    if not mcp_registry:
        raise HTTPException(status_code=501, detail="MCP tidak didukung")
    try:
        client = await mcp_registry.get_client(name)
        tools = await client.list_tools()
        return {"tools": tools}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class CallMcpToolRequest(BaseModel):
    tool_name: str
    arguments: dict

@app.post("/api/mcp/servers/{name}/call")
async def call_mcp_server_tool(name: str, req: CallMcpToolRequest):
    if not mcp_registry:
        raise HTTPException(status_code=501, detail="MCP tidak didukung")
    try:
        client = await mcp_registry.get_client(name)
        result = await client.call_tool(req.tool_name, req.arguments)
        return {"result": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
